#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Referências bibliográficas (APA 7) — inventário, extracção e formatação.

Princípio: nunca inventar metadados. Campo preenchido exige ``evidencia_*``
≠ ``vazio``. A Fase 0 só inventaria obras; a Fase 1 extrai com rastreio de
fonte; a formatação APA (Fase 3) opera só sobre TSV revisto.
"""

from __future__ import annotations

import json
import re
import unicodedata
import urllib.error
import urllib.parse
import urllib.request
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote, urlparse

import pandas as pd

# Prefixo típico do corpus local (CI não o possui — usar ``--raiz-corpus``).
PREFIXOS_CORPUS_OMISSAO: tuple[str, ...] = (
    r"E:\todos os textos",
    r"E:/todos os textos",
    "file:///E:/todos%20os%20textos",
    "file:///E:/todos os textos",
    "file://E:/todos%20os%20textos",
)

COLUNAS_INVENTARIO: tuple[str, ...] = (
    "doc_id",
    "caminho_ficheiro",
    "n_hits",
    "n_hits_nucleares",
    "ficheiro_existe",
    "tipo_provavel",
)

_TIPOS = (
    "grove",
    "tese",
    "actas",
    "artigo",
    "livro",
    "desconhecido",
)

# Nomes de ficheiro usam «_»; não exigir \b (underscore é carácter de palavra).
_RX_GROVE = re.compile(
    r"grove\s+music\s+online|oxford\s+music\s+online", re.I)
_RX_TESE = re.compile(
    r"thesis|dissertation|doutoramento|doutorado|(?<![a-z])tese(?![a-z])|"
    r"ph\.?\s*d\.?|m\.?\s*phil|mestrado|master'?s\s+thesis",
    re.I,
)
_RX_ACTAS = re.compile(
    r"proceedings|actas|atas|conference|symposium|congresso|"
    r"colloquium|workshop",
    re.I,
)
_RX_ARTIGO = re.compile(
    r"journal|revista|article|artigo|vol\.?\s*\d|pp\.?\s*\d",
    re.I,
)
_RX_LIVRO = re.compile(
    r"oxford\s+university\s+press|cambridge\s+university\s+press|"
    r"routledge|springer|handbook|edited\s+by",
    re.I,
)


def _truthy_nuclear(val) -> bool:
    if isinstance(val, bool):
        return val
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return False
    return str(val).strip().lower() in {
        "1", "true", "verdadeiro", "yes", "sim", "t", "y",
    }


def _strip_file_url(caminho: str) -> str:
    s = str(caminho or "").strip()
    if not s:
        return ""
    if s.lower().startswith("file:"):
        parsed = urlparse(s)
        path = unquote(parsed.path or "")
        # Windows: file:///E:/foo → /E:/foo
        if re.match(r"^/[A-Za-z]:", path):
            path = path[1:]
        return path.replace("/", "\\") if re.match(r"^[A-Za-z]:", path) else path
    return s


def remapear_caminho(
    caminho: str,
    raiz_corpus: Path | None = None,
    prefixos_origem: Iterable[str] | None = None,
) -> Path:
    """Resolve o caminho local, opcionalmente remapeando o prefixo do corpus.

    Se ``raiz_corpus`` for dada e o caminho (após descodificar ``file://``)
    começar por um dos ``prefixos_origem``, a parte relativa é juntada a
    ``raiz_corpus``. Caso contrário devolve o path tal qual (normalizado).
    """
    bruto = _strip_file_url(caminho)
    if not bruto:
        return Path()
    p = Path(bruto)
    if raiz_corpus is None:
        return p
    raiz = Path(raiz_corpus)
    prefixos = tuple(prefixos_origem) if prefixos_origem else PREFIXOS_CORPUS_OMISSAO
    # Comparação casefold + barras normalizadas
    norm = bruto.replace("\\", "/").casefold()
    for pref in prefixos:
        pref_s = _strip_file_url(pref).replace("\\", "/").rstrip("/")
        pref_n = pref_s.casefold()
        if norm == pref_n or norm.startswith(pref_n + "/"):
            rel = bruto.replace("\\", "/")[len(pref_s):].lstrip("/\\")
            return (raiz / rel) if rel else raiz
    return p


def ler_amostra_pdf(caminho: Path, *, max_paginas: int = 1,
                    max_chars: int = 2500) -> str:
    """Texto das primeiras páginas (vazio se ficheiro ausente / ilegível)."""
    if not caminho.is_file():
        return ""
    try:
        import fitz
    except ImportError:
        return ""
    try:
        doc = fitz.open(caminho)
    except Exception:
        return ""
    try:
        partes: list[str] = []
        n = min(max_paginas, doc.page_count)
        for i in range(n):
            try:
                partes.append(doc.load_page(i).get_text("text") or "")
            except Exception:
                continue
        return "\n".join(partes)[:max_chars]
    finally:
        doc.close()


def tipo_provavel(caminho: str, texto_pagina1: str = "") -> str:
    """Heurística *indicativa* — não é classificação bibliográfica final."""
    nome = Path(_strip_file_url(caminho)).name
    blob = f"{nome}\n{texto_pagina1}"
    if _RX_GROVE.search(blob):
        return "grove"
    if _RX_TESE.search(blob):
        return "tese"
    if _RX_ACTAS.search(blob):
        return "actas"
    if _RX_ARTIGO.search(blob):
        return "artigo"
    if _RX_LIVRO.search(blob):
        return "livro"
    return "desconhecido"


def _caminho_representativo(serie: pd.Series) -> str:
    vals = [
        str(v).strip() for v in serie
        if v is not None and not (isinstance(v, float) and pd.isna(v))
        and str(v).strip() and str(v).strip().lower() != "nan"
    ]
    if not vals:
        return ""
    # Mais frequente; empate → ordem lexicográfica estável
    contagem: dict[str, int] = {}
    for v in vals:
        contagem[v] = contagem.get(v, 0) + 1
    return sorted(contagem.items(), key=lambda kv: (-kv[1], kv[0]))[0][0]


def chave_doc_id(row: pd.Series) -> str:
    """Chave de agrupamento estável (doc_id ou fallback pelo nome do ficheiro)."""
    d = row.get("doc_id")
    if d is not None and not (isinstance(d, float) and pd.isna(d)):
        s = str(d).strip()
        if s and s.lower() != "nan":
            return s
    c = row.get("caminho_ficheiro", "")
    nome = Path(_strip_file_url(str(c))).name if c is not None else ""
    return f"sem_doc_id::{nome}" if nome else "sem_doc_id::"


def chaves_doc_id(df: pd.DataFrame) -> list[str]:
    return [chave_doc_id(r) for _, r in df.iterrows()]


def construir_inventario(
    df: pd.DataFrame,
    *,
    raiz_corpus: Path | None = None,
    prefixos_origem: Iterable[str] | None = None,
    ler_pdf: bool = True,
) -> pd.DataFrame:
    """Agrupa ``8_Concordancia`` por ``doc_id`` → linhas de inventário."""
    if "doc_id" not in df.columns:
        raise ValueError(
            "Folha sem coluna 'doc_id'. Corra textura_near antes do inventário."
        )
    work = df.copy()
    work["doc_id"] = [chave_doc_id(r) for _, r in work.iterrows()]

    col_path = "caminho_ficheiro" if "caminho_ficheiro" in work.columns else None
    nuc = (
        work["nuclear"].map(_truthy_nuclear)
        if "nuclear" in work.columns
        else pd.Series(False, index=work.index)
    )
    work = work.assign(_nuclear=nuc)

    rows: list[dict] = []
    for doc_id, grp in work.groupby("doc_id", sort=True):
        caminho = _caminho_representativo(grp[col_path]) if col_path else ""
        local = remapear_caminho(
            caminho, raiz_corpus=raiz_corpus, prefixos_origem=prefixos_origem)
        existe = bool(local.is_file()) if str(local) else False
        amostra = ler_amostra_pdf(local) if (ler_pdf and existe) else ""
        rows.append({
            "doc_id": str(doc_id),
            "caminho_ficheiro": caminho,
            "n_hits": int(len(grp)),
            "n_hits_nucleares": int(grp["_nuclear"].sum()),
            "ficheiro_existe": "sim" if existe else "nao",
            "tipo_provavel": tipo_provavel(caminho, amostra),
        })
    out = pd.DataFrame(rows, columns=list(COLUNAS_INVENTARIO))
    return out.reset_index(drop=True)


def escrever_inventario_tsv(df: pd.DataFrame, saida: Path) -> Path:
    saida = Path(saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(saida, sep="\t", index=False, encoding="utf-8", lineterminator="\n")
    return saida


def relatorio_inventario(df: pd.DataFrame, n_doc_excel: int) -> str:
    por_tipo = df["tipo_provavel"].value_counts().to_dict()
    n_miss = int((df["ficheiro_existe"] == "nao").sum())
    n_ok = int((df["ficheiro_existe"] == "sim").sum())
    linhas = [
        f"doc_id únicos (inventário): {len(df)}",
        f"doc_id únicos (Excel nunique): {n_doc_excel}",
        f"coincide: {'sim' if len(df) == n_doc_excel else 'NAO'}",
        f"ficheiros encontrados: {n_ok}",
        f"ficheiros em falta: {n_miss}",
        "tipo_provavel (indicativo):",
    ]
    for t in _TIPOS:
        linhas.append(f"  {t}: {int(por_tipo.get(t, 0))}")
    extras = sorted(set(por_tipo) - set(_TIPOS))
    for t in extras:
        linhas.append(f"  {t}: {int(por_tipo[t])}")
    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Fase 1 — extracção com evidência
# ---------------------------------------------------------------------------

CAMPOS_BIBLIO: tuple[str, ...] = (
    "autores",
    "ano",
    "titulo",
    "tipo",
    "contentor",
    "volume",
    "numero",
    "paginas",
    "editora",
    "doi_ou_url",
    "edicao",
)

EVIDENCIAS_OK: frozenset[str] = frozenset({
    "pdf_meta", "pagina1", "rodape", "nome_ficheiro", "manifesto",
    "doi_org", "vazio",
})

TIPOS_BIBLIO: frozenset[str] = frozenset({
    "artigo", "livro", "capitulo", "actas", "tese", "verbete",
    "relatorio", "desconhecido",
})

COLUNAS_RASCUNHO: tuple[str, ...] = (
    "doc_id",
    "caminho_ficheiro",
    *CAMPOS_BIBLIO,
    *[f"evidencia_{c}" for c in CAMPOS_BIBLIO],
    "confianca",
    "verificar",
    "caso_titulo",
    "duplicado_de",
)

_RX_DOI = re.compile(
    r"\b(10\.\d{4,9}/[-._;()/:A-Z0-9]+)\b", re.I)
_RX_DOI_URL = re.compile(
    r"https?://(?:dx\.)?doi\.org/(10\.\d{4,9}/[-._;()/:A-Z0-9]+)", re.I)
_RX_ANO = re.compile(r"\b((?:19|20)\d{2})\b")
_RX_FICHEIRO_ANO_TITULO = re.compile(
    r"^\((?P<ano>(?:19|20)\d{2})\)_(?P<titulo>.+?)$", re.I)
_RX_VOL = re.compile(
    r"\b(?:vol\.?|volume)\s*([0-9]+[A-Za-z]?)\b", re.I)
_RX_NUM = re.compile(
    r"\b(?:no\.?|n[ºo]\.?|number|issue)\s*([0-9]+)\b", re.I)
_RX_PP = re.compile(
    r"\bpp?\.\s*([0-9]+\s*[-–—]\s*[0-9]+)\b", re.I)
_RX_BY_AUTHOR = re.compile(
    r"(?im)^(?:by|por)\s+([A-ZÀ-Ý][\wÀ-ÿ'''\-]+(?:,\s*[A-ZÀ-Ý][\wÀ-ÿ.]+)?|"
    r"[A-ZÀ-Ý][\wÀ-ÿ'''\-]+(?:\s+[A-ZÀ-Ý]\.)+\s*[A-ZÀ-Ý][\wÀ-ÿ'''\-]+)\s*$")
_RX_AUTHOR_LINE = re.compile(
    r"(?m)^([A-ZÀ-Ý][\wÀ-ÿ'''\-]+,\s*[A-ZÀ-Ý](?:\.\s*[A-ZÀ-Ý]\.)*)\s*$")
_RX_GROVE_PUBLISHED = re.compile(
    r"(?:article\s+published\s+online|published\s+online|last\s+updated|"
    r"revised)\s*:?\s*"
    r"(?:(\d{1,2})\s+)?"
    r"(?:January|February|March|April|May|June|July|August|September|"
    r"October|November|December|[A-Za-zç]+)\s+"
    r"((?:19|20)\d{2})",
    re.I,
)
_RX_JOURNAL_LINE = re.compile(
    r"(?m)^(?P<journal>[A-ZÀ-Ý][^,\n]{2,80}?),\s*"
    r"(?:Vol\.?\s*(?P<vol>\d+)(?:,\s*)?)?"
    r"(?:No\.?\s*(?P<num>\d+)(?:,\s*)?)?"
    r"(?:pp?\.?\s*(?P<pp>\d+\s*[-–—]\s*\d+)(?:,\s*)?)?"
    r"(?P<ano>(?:19|20)\d{2})?\s*$",
    re.I,
)


@dataclass
class CampoEvidencia:
    valor: str = ""
    evidencia: str = "vazio"

    def preencher(self, valor: str, evidencia: str, *, forcar: bool = False) -> None:
        v = (valor or "").strip()
        if not v:
            return
        if evidencia not in EVIDENCIAS_OK or evidencia == "vazio":
            raise ValueError(f"evidência inválida: {evidencia!r}")
        if self.valor and not forcar:
            return
        self.valor = v
        self.evidencia = evidencia


@dataclass
class RascunhoObra:
    doc_id: str
    caminho_ficheiro: str = ""
    campos: dict[str, CampoEvidencia] = field(default_factory=dict)
    caso_titulo: str = ""
    duplicado_de: str = ""

    def __post_init__(self) -> None:
        for c in CAMPOS_BIBLIO:
            self.campos.setdefault(c, CampoEvidencia())

    def set(self, campo: str, valor: str, evidencia: str, *, forcar: bool = False) -> None:
        if campo not in self.campos:
            raise KeyError(campo)
        self.campos[campo].preencher(valor, evidencia, forcar=forcar)

    def get(self, campo: str) -> str:
        return self.campos[campo].valor

    def para_dict(self) -> dict[str, str]:
        out: dict[str, str] = {
            "doc_id": self.doc_id,
            "caminho_ficheiro": self.caminho_ficheiro,
        }
        for c in CAMPOS_BIBLIO:
            out[c] = self.campos[c].valor
            out[f"evidencia_{c}"] = self.campos[c].evidencia
        conf, verificar = _avaliar_confianca(self)
        out["confianca"] = conf
        out["verificar"] = verificar
        out["caso_titulo"] = self.caso_titulo or _caso_titulo(self.get("titulo"))
        out["duplicado_de"] = self.duplicado_de
        return out


def _norm_cmp(s: str) -> str:
    s = re.sub(r"\s+", " ", (s or "").strip().casefold())
    return re.sub(r"[^\w\s]", "", s, flags=re.UNICODE)


def _meta_coerente_com_pagina(meta_val: str, pagina: str) -> bool:
    """Só aceitar pdf_meta se o valor (ou um token longo) ocorrer na 1.ª página."""
    m = _norm_cmp(meta_val)
    p = _norm_cmp(pagina)
    if not m or not p:
        return False
    if m in p:
        return True
    tokens = [t for t in m.split() if len(t) >= 5]
    if not tokens:
        return False
    return sum(1 for t in tokens if t in p) >= max(1, len(tokens) // 2)


def _parece_lixo_meta(val: str) -> bool:
    v = (val or "").strip()
    if not v:
        return True
    lixo = (
        "microsoft word", "document1", "untitled", "scanner", "adobe pdf",
        "full text", "downloaded", "admin", "user",
    )
    vl = v.casefold()
    return any(x in vl for x in lixo) or len(v) < 3


def _parece_linha_titulo(ln: str) -> bool:
    if len(ln) < 8 or len(ln) > 140:
        return False
    if ln.lower().startswith("http"):
        return False
    # prosa corrida de scan
    if ln.endswith(".") and len(ln) > 70:
        return False
    return True


def _parece_title_case(ln: str) -> bool:
    palavras = [w for w in re.findall(r"[A-Za-zÀ-ÿ]{2,}", ln)]
    if len(palavras) < 2:
        return len(ln) <= 48
    maius = sum(1 for w in palavras if w[0].isupper())
    return maius >= max(2, (len(palavras) + 1) // 2)


def _caso_titulo(titulo: str) -> str:
    """Preservar capitalização; assinalar title case aparente para revisão humana."""
    t = (titulo or "").strip()
    if not t:
        return ""
    palavras = [w for w in re.findall(r"[A-Za-zÀ-ÿ]{3,}", t) if w.casefold() not in {
        "and", "the", "for", "with", "from", "into", "that", "this",
    }]
    if len(palavras) < 3:
        return ""
    maiusculas = sum(1 for w in palavras if w[0].isupper())
    if maiusculas / len(palavras) >= 0.7:
        return "por_rever"
    return "ok"


def _avaliar_confianca(r: RascunhoObra) -> tuple[str, str]:
    preenchidos = [
        c for c in CAMPOS_BIBLIO
        if r.campos[c].valor and r.campos[c].evidencia != "vazio"
    ]
    fortes = [
        c for c in preenchidos
        if r.campos[c].evidencia in {"pagina1", "rodape", "manifesto", "doi_org"}
    ]
    if r.get("tipo") == "desconhecido" or not r.get("titulo"):
        return "baixa", "sim"
    if len(fortes) >= 3 and r.get("ano"):
        return "alta", "nao" if r.get("autores") else "sim"
    if len(preenchidos) >= 2:
        return "media", "sim"
    return "baixa", "sim"


def _extrair_doi(texto: str) -> str:
    m = _RX_DOI_URL.search(texto or "")
    if m:
        return m.group(1)
    m = _RX_DOI.search(texto or "")
    return m.group(1) if m else ""


def ler_pdf_estruturado(caminho: Path, *, max_paginas: int = 2) -> dict:
    """Metadados + texto de corpo / cabeçalho / rodapé das primeiras páginas."""
    vazio = {
        "meta": {}, "pagina1": "", "rodape": "", "cabecalho": "", "existe": False,
    }
    if not caminho.is_file():
        return vazio
    try:
        import fitz
    except ImportError:
        return vazio
    try:
        doc = fitz.open(caminho)
    except Exception:
        return vazio
    try:
        meta_raw = doc.metadata or {}
        meta = {
            "title": str(meta_raw.get("title") or "").strip(),
            "author": str(meta_raw.get("author") or "").strip(),
            "subject": str(meta_raw.get("subject") or "").strip(),
        }
        corpos: list[str] = []
        rods: list[str] = []
        cabs: list[str] = []
        n = min(max_paginas, doc.page_count)
        for i in range(n):
            page = doc.load_page(i)
            h = float(page.rect.height) or 1.0
            for block in page.get_text("blocks") or []:
                if len(block) < 5:
                    continue
                y0, txt = float(block[1]), str(block[4] or "").strip()
                if not txt:
                    continue
                if y0 < h * 0.12:
                    cabs.append(txt)
                elif y0 > h * 0.85:
                    rods.append(txt)
                else:
                    corpos.append(txt)
            # fallback se blocks falharem
            if not corpos and not rods:
                corpos.append(page.get_text("text") or "")
        return {
            "meta": meta,
            # ordem de leitura: cabeçalho → corpo (Grove costuma estar no topo)
            "pagina1": "\n".join(cabs + corpos)[:6000],
            "rodape": "\n".join(rods)[:3000],
            "cabecalho": "\n".join(cabs)[:2000],
            "existe": True,
        }
    finally:
        doc.close()


def _aplicar_nome_ficheiro(r: RascunhoObra, *, grove: bool) -> None:
    stem = Path(_strip_file_url(r.caminho_ficheiro)).stem
    m = _RX_FICHEIRO_ANO_TITULO.match(stem)
    if not m:
        return
    # Grove: não usar ano do download/nome
    if not grove:
        r.set("ano", m.group("ano"), "nome_ficheiro")
    titulo = m.group("titulo").replace("_", " ").strip()
    if titulo:
        r.set("titulo", titulo, "nome_ficheiro")


def _aplicar_pdf_meta(r: RascunhoObra, meta: dict, pagina1: str) -> None:
    title = meta.get("title") or ""
    author = meta.get("author") or ""
    if title and not _parece_lixo_meta(title) and _meta_coerente_com_pagina(title, pagina1):
        r.set("titulo", title, "pdf_meta")
    if author and not _parece_lixo_meta(author) and _meta_coerente_com_pagina(author, pagina1):
        # normalizar "Jane Q. Smith" → manter; "Smith, J. Q." já OK
        r.set("autores", _normalizar_autor_meta(author), "pdf_meta")


def _normalizar_autor_meta(author: str) -> str:
    a = re.sub(r"\s+", " ", (author or "").strip())
    if not a:
        return ""
    if ";" in a:
        return a
    if re.match(r"^[\wÀ-ÿ'''\-]+,\s*", a):
        return a
    partes = a.split()
    if len(partes) >= 2 and "," not in a:
        apelido = partes[-1].rstrip(",")
        inits = " ".join(
            (f"{p[0]}." if p and p[0].isupper() else p) for p in partes[:-1]
        )
        return f"{apelido}, {inits}".replace("..", ".")
    return a


def _aplicar_grove(r: RascunhoObra, pagina1: str, rodape: str) -> bool:
    blob = f"{pagina1}\n{rodape}"
    if not _RX_GROVE.search(blob) and not _RX_GROVE.search(r.caminho_ficheiro):
        return False
    r.set("tipo", "verbete", "pagina1")
    r.set("contentor", "Grove Music Online", "pagina1")
    r.set("editora", "Oxford University Press", "pagina1")
    linhas = [ln.strip() for ln in pagina1.splitlines() if ln.strip()]
    grove_idxs = [i for i, ln in enumerate(linhas) if _RX_GROVE.search(ln)]

    def _candidato_titulo(ln: str) -> bool:
        if len(ln) < 3 or len(ln) > 120:
            return False
        low = ln.casefold()
        if low.startswith("http") or "doi.org" in low:
            return False
        if "published online" in low or "last updated" in low:
            return False
        if _RX_GROVE.search(ln) or _RX_BY_AUTHOR.match(ln):
            return False
        return True

    cands = [(i, ln) for i, ln in enumerate(linhas) if _candidato_titulo(ln)]
    if grove_idxs:
        g0 = grove_idxs[0]
        after = [ln for i, ln in cands if i > g0]
        before = [ln for i, ln in cands if i < g0]
        if after:
            r.set("titulo", after[0], "pagina1")
        elif before:
            r.set("titulo", before[0], "pagina1")
    elif cands:
        r.set("titulo", cands[0][1], "pagina1")

    for ln in linhas:
        m = _RX_BY_AUTHOR.match(ln)
        if m:
            r.set("autores", _normalizar_autor_meta(m.group(1)), "pagina1")
            break
    if not r.get("autores"):
        for ln in linhas:
            m2 = _RX_AUTHOR_LINE.match(ln)
            if not m2 or "grove" in ln.casefold():
                continue
            if r.get("titulo") and _norm_cmp(m2.group(1)) == _norm_cmp(r.get("titulo")):
                continue
            r.set("autores", m2.group(1), "pagina1")
            break
    mpub = _RX_GROVE_PUBLISHED.search(blob)
    if mpub:
        r.set("ano", mpub.group(2), "pagina1")
    doi = _extrair_doi(blob)
    if doi:
        r.set("doi_ou_url", f"https://doi.org/{doi}", "pagina1")
    return True


def _aplicar_pagina_e_rodape(r: RascunhoObra, pagina1: str, rodape: str) -> None:
    blob_rod = rodape or ""
    blob_pag = pagina1 or ""
    doi = _extrair_doi(blob_rod) or _extrair_doi(blob_pag)
    if doi:
        fonte = "rodape" if _extrair_doi(blob_rod) else "pagina1"
        r.set("doi_ou_url", f"https://doi.org/{doi}", fonte)

    for fonte, texto in (("rodape", blob_rod), ("pagina1", blob_pag)):
        if not texto:
            continue
        mj = _RX_JOURNAL_LINE.search(texto)
        if mj:
            if mj.group("journal"):
                r.set("contentor", mj.group("journal").strip(), fonte)
                if not r.get("tipo") or r.get("tipo") == "desconhecido":
                    r.set("tipo", "artigo", fonte)
            if mj.group("vol"):
                r.set("volume", mj.group("vol"), fonte)
            if mj.group("num"):
                r.set("numero", mj.group("num"), fonte)
            if mj.group("pp"):
                r.set("paginas", mj.group("pp").replace("—", "-").replace("–", "-"),
                      fonte)
            if mj.group("ano"):
                r.set("ano", mj.group("ano"), fonte)
        else:
            mvol, mnum, mpp, mano = (
                _RX_VOL.search(texto), _RX_NUM.search(texto),
                _RX_PP.search(texto), _RX_ANO.search(texto),
            )
            if mvol:
                r.set("volume", mvol.group(1), fonte)
            if mnum:
                r.set("numero", mnum.group(1), fonte)
            if mpp:
                r.set("paginas", mpp.group(1).replace("—", "-").replace("–", "-"),
                      fonte)
            if mano and fonte == "rodape":
                r.set("ano", mano.group(1), fonte)

    # título: só linhas que parecem título (não prosa de página digitalizada)
    if not r.get("titulo"):
        sinais = bool(r.get("contentor") or r.get("doi_ou_url") or r.get("volume"))
        for ln in blob_pag.splitlines():
            ln = ln.strip()
            if not _parece_linha_titulo(ln):
                continue
            if _RX_BY_AUTHOR.match(ln) or _RX_AUTHOR_LINE.match(ln):
                continue
            if _RX_GROVE.search(ln) or _RX_JOURNAL_LINE.match(ln):
                continue
            if sinais or _parece_title_case(ln):
                r.set("titulo", ln, "pagina1")
                break

    if not r.get("autores"):
        for ln in blob_pag.splitlines():
            ln = ln.strip()
            m = _RX_BY_AUTHOR.match(ln)
            if m:
                r.set("autores", _normalizar_autor_meta(m.group(1)), "pagina1")
                break
            m2 = _RX_AUTHOR_LINE.match(ln)
            if m2:
                r.set("autores", m2.group(1), "pagina1")
                break
            # «Jane Q. Smith» numa linha só
            if re.match(
                r"^[A-ZÀ-Ý][\wÀ-ÿ'''\-]+(?:\s+[A-ZÀ-Ý]\.)+\s+[A-ZÀ-Ý][\wÀ-ÿ'''\-]+$",
                ln,
            ):
                r.set("autores", _normalizar_autor_meta(ln), "pagina1")
                break


def _classificar_tipo(r: RascunhoObra, caminho: str, pagina1: str) -> None:
    if r.get("tipo") and r.get("tipo") in TIPOS_BIBLIO and r.get("tipo") != "desconhecido":
        return
    blob = f"{Path(caminho).name}\n{pagina1}"
    if _RX_GROVE.search(blob):
        r.set("tipo", "verbete", "pagina1" if _RX_GROVE.search(pagina1) else "nome_ficheiro")
        return
    if _RX_TESE.search(blob):
        r.set("tipo", "tese",
              "pagina1" if _RX_TESE.search(pagina1) else "nome_ficheiro")
        return
    if _RX_ACTAS.search(blob):
        r.set("tipo", "actas",
              "pagina1" if _RX_ACTAS.search(pagina1) else "nome_ficheiro")
        return
    if r.get("contentor") or _RX_ARTIGO.search(blob):
        r.set("tipo", "artigo",
              "rodape" if r.campos["contentor"].evidencia == "rodape"
              else ("pagina1" if _RX_ARTIGO.search(pagina1) else "nome_ficheiro"))
        return
    if _RX_LIVRO.search(blob):
        r.set("tipo", "livro",
              "pagina1" if _RX_LIVRO.search(pagina1) else "nome_ficheiro")
        return
    r.campos["tipo"].valor = "desconhecido"
    r.campos["tipo"].evidencia = "nome_ficheiro"


def ler_manifesto(xlsx: Path | None) -> dict[str, dict[str, str]]:
    """Mapa path-normalizado → campos úteis do Manifesto_corpus (se existirem)."""
    if xlsx is None or not Path(xlsx).is_file():
        return {}
    try:
        raw = pd.read_excel(xlsx, sheet_name="Manifesto_corpus")
    except Exception:
        return {}
    if raw is None or raw.empty:
        return {}
    cols = {c.casefold(): c for c in raw.columns}
    col_path = next(
        (cols[k] for k in ("caminho", "caminho_ficheiro", "path", "ficheiro")
         if k in cols),
        None,
    )
    if col_path is None:
        return {}
    aliases = {
        "autores": ("autor", "autores", "author", "authors"),
        "ano": ("ano", "year", "data"),
        "titulo": ("titulo", "title", "obra"),
        "tipo": ("tipo", "type"),
        "doi_ou_url": ("doi", "url", "doi_ou_url"),
        "editora": ("editora", "publisher"),
        "contentor": ("contentor", "revista", "journal"),
    }
    out: dict[str, dict[str, str]] = {}
    for _, row in raw.iterrows():
        chave = _norm_chave_path(str(row.get(col_path) or ""))
        if not chave:
            continue
        campos: dict[str, str] = {}
        for dest, cands in aliases.items():
            for cand in cands:
                if cand in cols:
                    val = row.get(cols[cand])
                    if val is not None and not (isinstance(val, float) and pd.isna(val)):
                        s = str(val).strip()
                        if s and s.lower() != "nan":
                            campos[dest] = s
                            break
        if campos:
            out[chave] = campos
    return out


def _norm_chave_path(s: str) -> str:
    s = _strip_file_url(s).replace("\\", "/").casefold()
    return Path(s).name


def _aplicar_manifesto(r: RascunhoObra, manifesto: dict[str, dict[str, str]]) -> None:
    chave = _norm_chave_path(r.caminho_ficheiro)
    dados = manifesto.get(chave) or {}
    for campo, val in dados.items():
        if campo in r.campos:
            r.set(campo, val, "manifesto")


def resolver_doi_org(doi: str, timeout: float = 12.0) -> dict[str, str]:
    """Resolve DOI via doi.org → Crossref (só com ``--permitir-web``)."""
    doi = (doi or "").strip()
    doi = re.sub(r"^https?://(?:dx\.)?doi\.org/", "", doi, flags=re.I)
    if not doi.startswith("10."):
        return {}
    url = f"https://api.crossref.org/works/{urllib.parse.quote(doi)}"
    req = urllib.request.Request(
        url,
        headers={"User-Agent": "TermStatistics-Referencias/1.0 (research)"},
    )
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            data = json.loads(resp.read().decode("utf-8", errors="replace"))
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError):
        return {}
    msg = (data or {}).get("message") or {}
    out: dict[str, str] = {}
    title = msg.get("title") or []
    if title:
        out["titulo"] = str(title[0]).strip()
    container = msg.get("container-title") or []
    if container:
        out["contentor"] = str(container[0]).strip()
    issued = (msg.get("issued") or {}).get("date-parts") or []
    if issued and issued[0]:
        out["ano"] = str(issued[0][0])
    vol = msg.get("volume")
    if vol:
        out["volume"] = str(vol)
    issue = msg.get("issue")
    if issue:
        out["numero"] = str(issue)
    page = msg.get("page")
    if page:
        out["paginas"] = str(page)
    publisher = msg.get("publisher")
    if publisher:
        out["editora"] = str(publisher)
    authors = msg.get("author") or []
    bits = []
    for a in authors:
        family = (a.get("family") or "").strip()
        given = (a.get("given") or "").strip()
        if family and given:
            inits = " ".join(p[0] + "." for p in given.replace("-", " ").split() if p)
            bits.append(f"{family}, {inits}")
        elif family:
            bits.append(family)
    if bits:
        out["autores"] = "; ".join(bits)
    tipo_cr = (msg.get("type") or "").lower()
    mapa = {
        "journal-article": "artigo",
        "book": "livro",
        "book-chapter": "capitulo",
        "proceedings-article": "actas",
        "dissertation": "tese",
        "report": "relatorio",
    }
    if tipo_cr in mapa:
        out["tipo"] = mapa[tipo_cr]
    return out


def _aplicar_doi_org(r: RascunhoObra) -> None:
    doi = r.get("doi_ou_url")
    if not doi:
        return
    meta = resolver_doi_org(doi)
    for campo, val in meta.items():
        if campo in r.campos and not r.get(campo):
            r.set(campo, val, "doi_org")


def extrair_obra(
    doc_id: str,
    caminho_ficheiro: str,
    *,
    raiz_corpus: Path | None = None,
    prefixos_origem: Iterable[str] | None = None,
    manifesto: dict[str, dict[str, str]] | None = None,
    permitir_web: bool = False,
) -> RascunhoObra:
    """Extrai um rascunho bibliográfico com evidência por campo."""
    r = RascunhoObra(doc_id=doc_id, caminho_ficheiro=caminho_ficheiro)
    local = remapear_caminho(
        caminho_ficheiro, raiz_corpus=raiz_corpus, prefixos_origem=prefixos_origem)
    pdf = ler_pdf_estruturado(local)
    pagina1, rodape = pdf["pagina1"], pdf["rodape"]
    grove = _aplicar_grove(r, pagina1, rodape)
    _aplicar_nome_ficheiro(r, grove=grove)
    if pdf["existe"]:
        _aplicar_pdf_meta(r, pdf["meta"], pagina1)
        if not grove:
            _aplicar_pagina_e_rodape(r, pagina1, rodape)
    if manifesto:
        _aplicar_manifesto(r, manifesto)
    _classificar_tipo(r, caminho_ficheiro, pagina1)
    if permitir_web:
        _aplicar_doi_org(r)
    # Garantir tipo sempre com evidência
    if not r.campos["tipo"].valor:
        r.campos["tipo"].valor = "desconhecido"
        r.campos["tipo"].evidencia = "nome_ficheiro"
    r.caso_titulo = _caso_titulo(r.get("titulo"))
    return r


def _chave_duplicado(r: RascunhoObra) -> str | None:
    a, y, t = r.get("autores"), r.get("ano"), r.get("titulo")
    if not (a and y and t):
        return None
    return f"{_norm_cmp(a)}|{y}|{_norm_cmp(t)}"


def marcar_duplicados(obras: list[RascunhoObra]) -> None:
    """Preenche ``duplicado_de`` quando autor+ano+título coincidem (canónico = 1.º)."""
    visto: dict[str, str] = {}
    for r in sorted(obras, key=lambda x: x.doc_id):
        k = _chave_duplicado(r)
        if not k:
            continue
        if k in visto:
            r.duplicado_de = visto[k]
        else:
            visto[k] = r.doc_id


def assert_evidencias_consistentes(df: pd.DataFrame) -> None:
    """Nenhum campo preenchido sem evidencia_* correspondente (≠ vazio)."""
    for i, row in df.iterrows():
        for c in CAMPOS_BIBLIO:
            val = str(row.get(c) or "").strip()
            ev = str(row.get(f"evidencia_{c}") or "vazio").strip()
            if val and ev == "vazio":
                raise AssertionError(
                    f"linha {i} doc_id={row.get('doc_id')}: "
                    f"campo {c!r}={val!r} sem evidência"
                )
            if ev not in EVIDENCIAS_OK:
                raise AssertionError(
                    f"linha {i}: evidência inválida {ev!r} em {c}"
                )


def construir_rascunho(
    df: pd.DataFrame,
    *,
    xlsx: Path | None = None,
    raiz_corpus: Path | None = None,
    prefixos_origem: Iterable[str] | None = None,
    permitir_web: bool = False,
) -> pd.DataFrame:
    """Uma linha por ``doc_id`` → ``referencias_rascunho.tsv``."""
    if "doc_id" not in df.columns:
        raise ValueError("Folha sem coluna doc_id.")
    work = df.copy()
    work["_did"] = [chave_doc_id(r) for _, r in work.iterrows()]
    col_path = "caminho_ficheiro" if "caminho_ficheiro" in work.columns else None
    manifesto = ler_manifesto(xlsx) if xlsx else {}
    obras: list[RascunhoObra] = []
    for doc_id, grp in work.groupby("_did", sort=True):
        caminho = _caminho_representativo(grp[col_path]) if col_path else ""
        obras.append(
            extrair_obra(
                str(doc_id),
                caminho,
                raiz_corpus=raiz_corpus,
                prefixos_origem=prefixos_origem,
                manifesto=manifesto,
                permitir_web=permitir_web,
            )
        )
    marcar_duplicados(obras)
    out = pd.DataFrame([o.para_dict() for o in obras], columns=list(COLUNAS_RASCUNHO))
    assert_evidencias_consistentes(out)
    return out.reset_index(drop=True)


def escrever_rascunho_tsv(df: pd.DataFrame, saida: Path) -> Path:
    saida = Path(saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(saida, sep="\t", index=False, encoding="utf-8", lineterminator="\n")
    return saida


def relatorio_rascunho(df: pd.DataFrame) -> str:
    n = len(df)
    n_ver = int((df["verificar"] == "sim").sum()) if n else 0
    por_tipo = df["tipo"].value_counts().to_dict() if n else {}
    n_dup = int((df["duplicado_de"].astype(str).str.strip() != "").sum()) if n else 0
    linhas = [
        f"obras (doc_id): {n}",
        f"verificar=sim: {n_ver}",
        f"verificar=nao: {n - n_ver}",
        f"duplicado_de preenchido: {n_dup}",
        "tipo:",
    ]
    for t in sorted(TIPOS_BIBLIO):
        if t in por_tipo:
            linhas.append(f"  {t}: {int(por_tipo[t])}")
    for t, v in sorted(por_tipo.items()):
        if t not in TIPOS_BIBLIO:
            linhas.append(f"  {t}: {int(v)}")
    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Fase 3 — formatação APA 7 + citação curta
# ---------------------------------------------------------------------------

COLUNAS_APA7: tuple[str, ...] = (
    "doc_id",
    "citacao_curta",
    "referencia_apa7",
    "tipo",
    "verificar",
    "completa",
)


def chave_ordenacao_apa(texto: str) -> str:
    """Chave alfabética sem diacríticos (apelidos PT/FR)."""
    s = unicodedata.normalize("NFD", (texto or "").strip().casefold())
    s = "".join(ch for ch in s if unicodedata.category(ch) != "Mn")
    return s


def _g(row: dict | pd.Series, key: str) -> str:
    v = row.get(key, "") if hasattr(row, "get") else ""
    if v is None or (isinstance(v, float) and pd.isna(v)):
        return ""
    s = str(v).strip()
    return "" if s.lower() == "nan" else s


def _autores_lista(autores: str) -> list[str]:
    """Parte 'Apelido, N.; Apelido, N.' → lista de unidades autor."""
    raw = (autores or "").strip()
    if not raw:
        return []
    partes = [p.strip() for p in re.split(r"\s*;\s*", raw) if p.strip()]
    return partes


def _apelido(autor_unit: str) -> str:
    a = (autor_unit or "").strip()
    if not a:
        return ""
    if "," in a:
        return a.split(",", 1)[0].strip()
    return a.split()[-1]


def _ano_cita(ano: str, sufixo: str = "") -> str:
    a = (ano or "").strip()
    if not a:
        return "n.d." + sufixo
    return a + sufixo


def citacao_curta_de(
    autores: str,
    ano: str,
    *,
    sufixo_ano: str = "",
) -> str:
    """Apelido, AAAA | Apelido & Apelido, AAAA | Apelido et al., AAAA."""
    units = _autores_lista(autores)
    y = _ano_cita(ano, sufixo_ano)
    if not units:
        return f"(Autor desconhecido, {y})"
    if len(units) == 1:
        return f"{_apelido(units[0])}, {y}"
    if len(units) == 2:
        return f"{_apelido(units[0])} & {_apelido(units[1])}, {y}"
    return f"{_apelido(units[0])} et al., {y}"


def _ital(s: str) -> str:
    s = (s or "").strip()
    return f"*{s}*" if s else ""


def _ponto(s: str) -> str:
    s = (s or "").rstrip()
    if not s:
        return ""
    if s.endswith((".", "!", "?", "/", "»", '"')):
        return s
    return s + "."


def _doi_terminal(doi_ou_url: str) -> str:
    d = (doi_ou_url or "").strip()
    if not d:
        return ""
    if d.lower().startswith("http"):
        return d
    if d.startswith("10."):
        return f"https://doi.org/{d}"
    return d


def formatar_referencia_apa7(row: dict | pd.Series) -> str:
    """Uma referência APA 7 (itálicos em ``*…*``). Não inventa campos em falta."""
    tipo = (_g(row, "tipo") or "desconhecido").lower()
    autores = _g(row, "autores") or "Autor desconhecido"
    ano = _g(row, "ano") or "n.d."
    titulo = _g(row, "titulo")
    contentor = _g(row, "contentor")
    volume = _g(row, "volume")
    numero = _g(row, "numero")
    paginas = _g(row, "paginas").replace("—", "-").replace("–", "-")
    editora = _g(row, "editora")
    edicao = _g(row, "edicao")
    doi = _doi_terminal(_g(row, "doi_ou_url"))

    def _tail_doi(base: str) -> str:
        base = base.rstrip()
        if doi:
            # APA 7: sem ponto após DOI/URL
            if not base.endswith("."):
                base = _ponto(base)
            return f"{base} {doi}"
        return _ponto(base)

    if tipo == "artigo":
        # Autor. (Ano). Título. *Revista, vol*(nº), pp–pp. DOI
        bits = [f"{autores} ({ano})."]
        if titulo:
            bits.append(_ponto(titulo))
        rev = contentor
        if rev and volume:
            mid = f"{_ital(f'{rev}, {volume}')}"
            if numero:
                mid += f"({numero})"
            if paginas:
                mid += f", {paginas}"
            bits.append(_ponto(mid) if not doi else mid + ".")
        elif rev:
            bits.append(_ponto(_ital(rev)) if not doi else _ital(rev) + ".")
        return _tail_doi(" ".join(bits))

    if tipo == "verbete":
        # Autor. (Ano). Título. In *Grove…*. Editora. DOI
        cont = contentor or "Grove Music Online"
        ed = editora or "Oxford University Press"
        bits = [f"{autores} ({ano})."]
        if titulo:
            bits.append(_ponto(titulo))
        bits.append(f"In {_ital(cont)}.")
        bits.append(_ponto(ed) if not doi else ed + ".")
        return _tail_doi(" ".join(bits))

    if tipo == "livro":
        bits = [f"{autores} ({ano})."]
        if titulo:
            tit = _ital(titulo)
            if edicao:
                tit = f"{tit} ({edicao})"
            bits.append(_ponto(tit) if not (editora or doi) else tit + ".")
        if editora:
            bits.append(_ponto(editora) if not doi else editora + ".")
        return _tail_doi(" ".join(bits))

    if tipo == "capitulo":
        bits = [f"{autores} ({ano})."]
        if titulo:
            bits.append(_ponto(titulo))
        if contentor:
            mid = f"In {_ital(contentor)}"
            if paginas:
                mid += f" (pp. {paginas})"
            bits.append(_ponto(mid) if not (editora or doi) else mid + ".")
        if editora:
            bits.append(_ponto(editora) if not doi else editora + ".")
        return _tail_doi(" ".join(bits))

    if tipo == "actas":
        bits = [f"{autores} ({ano})."]
        if titulo:
            bits.append(_ponto(titulo))
        if contentor:
            mid = f"In {_ital(contentor)}"
            if paginas:
                mid += f" (pp. {paginas})"
            bits.append(_ponto(mid) if not (editora or doi) else mid + ".")
        if editora:
            bits.append(_ponto(editora) if not doi else editora + ".")
        return _tail_doi(" ".join(bits))

    if tipo == "tese":
        bits = [f"{autores} ({ano})."]
        if titulo:
            inst = editora  # instituição em editora (contrato Fase 2)
            tag = f" [Doctoral dissertation, {inst}]" if inst else " [Doctoral dissertation]"
            bits.append(_ital(titulo) + tag + ".")
        elif editora:
            bits.append(f"[Doctoral dissertation, {editora}].")
        return _tail_doi(" ".join(bits))

    if tipo == "relatorio":
        bits = [f"{autores} ({ano})."]
        if titulo:
            bits.append(_ponto(_ital(titulo)) if not (editora or doi) else _ital(titulo) + ".")
        if editora:
            bits.append(_ponto(editora) if not doi else editora + ".")
        return _tail_doi(" ".join(bits))

    # desconhecido — só o que existir, sem fabricar contentor/editora
    bits = [f"{autores} ({ano})."]
    if titulo:
        bits.append(_ponto(titulo) if not doi else titulo + ".")
    elif _g(row, "caminho_ficheiro"):
        bits.append(_ponto(Path(_strip_file_url(_g(row, "caminho_ficheiro"))).stem.replace("_", " ")))
    return _tail_doi(" ".join(bits))


def referencia_completa_minima(row: dict | pd.Series) -> bool:
    """Mínimo para listar na bibliografia 'completa': autores+título+tipo conhecido, ou verbete com contentor."""
    tipo = (_g(row, "tipo") or "desconhecido").lower()
    if tipo == "desconhecido":
        return False
    if not _g(row, "titulo"):
        return False
    if tipo == "verbete":
        return bool(_g(row, "contentor") or _g(row, "editora"))
    if tipo == "artigo":
        return bool(_g(row, "contentor") or _g(row, "doi_ou_url"))
    return bool(_g(row, "autores"))


def desambiguar_anos(df: pd.DataFrame) -> pd.DataFrame:
    """Sufixos a/b/c determinísticos (ordenados por título) para mesmo apelido+ano."""
    out = df.copy()
    out["_apelido0"] = [
        _apelido(_autores_lista(_g(r, "autores"))[0]) if _autores_lista(_g(r, "autores")) else ""
        for _, r in out.iterrows()
    ]
    out["_ano_base"] = [
        _g(r, "ano") if _g(r, "ano") else "n.d." for _, r in out.iterrows()
    ]
    out["_tit_ord"] = out.apply(
        lambda r: chave_ordenacao_apa(_g(r, "titulo") or _g(r, "caminho_ficheiro")),
        axis=1,
    )
    out["_sufixo"] = ""
    for (_, _), grp in out.groupby(["_apelido0", "_ano_base"], sort=False):
        if len(grp) < 2:
            continue
        if not grp.iloc[0]["_apelido0"] and grp.iloc[0]["_ano_base"] == "n.d.":
            continue
        ordenados = grp.sort_values(["_tit_ord", "doc_id"]).index.tolist()
        if len(ordenados) == 1:
            continue
        for i, idx in enumerate(ordenados):
            out.at[idx, "_sufixo"] = chr(ord("a") + i)
    return out


def construir_apa7(df_rascunho: pd.DataFrame) -> pd.DataFrame:
    """A partir do TSV rascunho/revisto → citacao_curta + referencia_apa7."""
    if df_rascunho is None or df_rascunho.empty:
        return pd.DataFrame(columns=list(COLUNAS_APA7))
    work = desambiguar_anos(df_rascunho)
    rows = []
    for _, r in work.iterrows():
        suf = str(r.get("_sufixo") or "")
        ano = _g(r, "ano")
        cita = citacao_curta_de(_g(r, "autores"), ano, sufixo_ano=suf)
        # inject suffix into APA year
        row_fmt = {c: _g(r, c) for c in list(CAMPOS_BIBLIO) + ["doc_id", "caminho_ficheiro", "verificar"]}
        if suf:
            row_fmt["ano"] = (ano or "n.d.") + suf
        apa = formatar_referencia_apa7(row_fmt)
        rows.append({
            "doc_id": _g(r, "doc_id"),
            "citacao_curta": cita,
            "referencia_apa7": apa,
            "tipo": _g(r, "tipo") or "desconhecido",
            "verificar": _g(r, "verificar") or "sim",
            "completa": "sim" if referencia_completa_minima(r) else "nao",
            "_ord_autor": chave_ordenacao_apa(
                _apelido(_autores_lista(_g(r, "autores"))[0])
                if _autores_lista(_g(r, "autores")) else _g(r, "titulo")
            ),
            "_ord_ano": _g(r, "ano") or "9999",
            "_ord_tit": chave_ordenacao_apa(_g(r, "titulo")),
        })
    out = pd.DataFrame(rows)
    out = out.sort_values(
        ["_ord_autor", "_ord_ano", "_ord_tit", "doc_id"], kind="mergesort"
    ).reset_index(drop=True)
    return out[list(COLUNAS_APA7)]


def escrever_apa7_tsv(df: pd.DataFrame, saida: Path) -> Path:
    saida = Path(saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    df.to_csv(saida, sep="\t", index=False, encoding="utf-8", lineterminator="\n")
    return saida


def escrever_apa7_md(df: pd.DataFrame, saida: Path, *, titulo: str = "Referências") -> Path:
    saida = Path(saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    linhas = [
        f"# {titulo}",
        "",
        "> Rascunho gerado automaticamente (extracção com evidência). "
        "Entradas incompletas exigem revisão humana antes de uso dissertativo.",
        "",
    ]
    completas = df[df["completa"] == "sim"] if "completa" in df.columns else df
    incompletas = df[df["completa"] != "sim"] if "completa" in df.columns else df.iloc[0:0]
    linhas.append(f"## Lista APA 7 ({len(completas)} entradas com campos mínimos)")
    linhas.append("")
    for _, r in completas.iterrows():
        linhas.append(f"- {r['referencia_apa7']}")
        linhas.append("")
    if len(incompletas):
        linhas.append(f"## Por rever ({len(incompletas)} — dados insuficientes)")
        linhas.append("")
        for _, r in incompletas.iterrows():
            linhas.append(
                f"- `{r['doc_id']}` — {r['citacao_curta']}: {r['referencia_apa7']}"
            )
            linhas.append("")
    saida.write_text("\n".join(linhas), encoding="utf-8")
    return saida


def escrever_apa7_docx(
    df: pd.DataFrame,
    saida: Path,
    *,
    titulo: str = "Referências",
    nota: str = "",
) -> Path:
    """DOCX autónomo com secção Referências (avanço pendente)."""
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Cm, Pt

    from textura_apendice import _add_text_with_italic_md

    saida = Path(saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2.5)

    h = doc.add_heading(titulo, level=1)
    for run in h.runs:
        run.font.name = "Times New Roman"

    nota = nota or (
        "Rascunho bibliográfico gerado por extracção automática com rastreio "
        "de evidência (Term_Statistics). Não inventa metadados: entradas sem "
        "campos mínimos ficam na secção «Por rever». A revisão humana do TSV "
        "é obrigatória antes de citar na dissertação."
    )
    pnota = doc.add_paragraph()
    _add_text_with_italic_md(pnota, nota, sz=19)
    for run in pnota.runs:
        run.italic = True

    completas = df[df["completa"] == "sim"] if "completa" in df.columns else df
    incompletas = (
        df[df["completa"] != "sim"] if "completa" in df.columns
        else df.iloc[0:0]
    )

    doc.add_heading(f"Lista APA 7.ª ed. ({len(completas)})", level=2)
    for _, r in completas.iterrows():
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Cm(-0.75)
        pf.left_indent = Cm(0.75)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        _add_text_with_italic_md(p, str(r["referencia_apa7"]), sz=19)

    if len(incompletas):
        doc.add_heading(f"Por rever ({len(incompletas)})", level=2)
        p0 = doc.add_paragraph()
        _add_text_with_italic_md(
            p0,
            "Entradas sem título/tipo/contentor suficientes — "
            "não usar como referência final.",
            sz=19,
        )
        for _, r in incompletas.iterrows():
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.first_line_indent = Cm(-0.75)
            pf.left_indent = Cm(0.75)
            pf.space_after = Pt(4)
            texto = (
                f"{r['citacao_curta']} — {r['referencia_apa7']} "
                f"[doc_id={r['doc_id']}]"
            )
            _add_text_with_italic_md(p, texto, sz=18)

    doc.save(str(saida))
    return saida


def relatorio_apa7(df: pd.DataFrame) -> str:
    n = len(df)
    n_ok = int((df["completa"] == "sim").sum()) if n and "completa" in df.columns else 0
    return "\n".join([
        f"referências formatadas: {n}",
        f"completas (mínimo bibliográfico): {n_ok}",
        f"por rever: {n - n_ok}",
    ])
