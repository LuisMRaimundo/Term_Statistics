# -*- coding: utf-8 -*-
"""Golden-master E2E: matriz_miniatura → near.xlsx → análise → DOCX.

Locks sheet names, ``8_Concordancia`` column order, nuclear counts,
duplicate-tag families, association OR for ``homogeneous``, and appendix
table row count. Any behavioural change in extraction/classification
must update these expectations with an explicit justification.
"""

from __future__ import annotations

import json
import os
import re
import subprocess
import sys
import time
from pathlib import Path

import pandas as pd
import pytest
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
FIXTURES = Path(__file__).resolve().parent / "fixtures"
MATRIZ = FIXTURES / "matriz_miniatura.xlsx"
CAMPO = FIXTURES / "campo_miniatura.txt"
BASELINE = FIXTURES / "golden_near_baseline.json"

SHEETS_NEAR = [
    "0_Instrucoes",
    "Config_lexico",
    "Manifesto_corpus",
    "8_Concordancia",
    "8_Concordancia_Hits",
    "8_Concordancia_Ocorrencias",
    "9_Excluidas",
    "Duplicados",
]

COLS_CONCORDANCIA = [
    "source_matrix_row", "texture_occurrence_id", "match_id", "hit_key",
    "grupo_passagem_id", "candidato_duplicado",
    "no", "termo_tipo", "canonical_term", "query_pattern",
    "termo_forma", "matched_form", "n_palavras", "distancia", "lado",
    "negado", "graduado", "modalizado", "relacao_sintactica",
    "polaridade_base", "polaridade", "eixo",
    "censurado_esq", "censurado_dir",
    "idx_no", "idx_termo", "off_no", "off_termo",
    "n_nos_janela", "forma_em_composto",
    "caminho_ficheiro", "doc_id", "url", "contexto",
    "motivo_exclusao", "nuclear", "fonte_classificacao",
    "n_janelas_fundidas", "revisao_sugerida",
    "nucleo_da_propriedade", "orientacao", "governante", "percurso_dep",
    "dominio", "dominio_janela", "revisto_por_humano", "nota_revisao",
]

# Captured under --lingua en (en_core_web_sm-3.8.0 + campo_miniatura).
# PT fixture rows are excluded by NOS language filter; counts drop vs todas.
N_HITS = 24
N_NUCLEAR = 20
N_NON_NUCLEAR = 4
OR_HOMOGENEOUS = 9.267
N_APENDICE_DATA_ROWS = 20

RX_C = re.compile(r"citacao_entre_doc_ids:C\d{4}")
RX_P = re.compile(r"passagem_sobreposta:P\d{4}")
RX_J = re.compile(r"janela_sobreposta:J\d{4}")


def _run(args: list[str], *, cwd: Path = ROOT) -> subprocess.CompletedProcess:
    return subprocess.run(
        [sys.executable, *args],
        cwd=str(cwd),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
    )


def _truthy_nuclear(s: pd.Series) -> pd.Series:
    return s.map(
        lambda v: v is True or str(v).strip().lower() in {"true", "1", "sim"}
    )


@pytest.fixture(scope="module")
def near_bundle(tmp_path_factory) -> dict:
    assert MATRIZ.is_file(), f"missing fixture {MATRIZ}"
    assert CAMPO.is_file(), f"missing fixture {CAMPO}"
    out_dir = tmp_path_factory.mktemp("golden_near")
    out = out_dir / "golden_near.xlsx"
    t0 = time.perf_counter()
    proc = _run([
        "textura_near.py",
        "--xlsx", str(MATRIZ),
        "--saida", str(out),
        "--near", "4",
        # EN-only: PT/FR fixture rows must not enter this golden.
        "--lingua", "en",
        "--termos", str(CAMPO),
        "--dominio-omissao", "musicologia",
    ])
    elapsed = time.perf_counter() - t0
    assert proc.returncode == 0, (
        f"textura_near failed ({proc.returncode})\n"
        f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
    )
    assert out.is_file()
    (out_dir / "near_seconds.txt").write_text(
        f"{elapsed:.3f}\n", encoding="utf-8")
    return {"path": out, "seconds": elapsed}


@pytest.fixture(scope="module")
def near_xlsx(near_bundle) -> Path:
    return near_bundle["path"]


@pytest.fixture(scope="module")
def analise_xlsx(near_xlsx: Path, tmp_path_factory) -> Path:
    out_dir = tmp_path_factory.mktemp("golden_analise")
    out = out_dir / "golden_analise.xlsx"
    proc = _run([
        "textura_analise.py",
        "--xlsx", str(near_xlsx),
        "--saida", str(out),
    ])
    assert proc.returncode == 0, (
        f"textura_analise failed ({proc.returncode})\n"
        f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
    )
    assert out.is_file()
    return out


@pytest.fixture(scope="module")
def apendice_docx(analise_xlsx: Path, tmp_path_factory) -> Path:
    out_dir = tmp_path_factory.mktemp("golden_apendice")
    out = out_dir / "golden_apendice.docx"
    proc = _run([
        "textura_apendice.py",
        "--xlsx", str(analise_xlsx),
        "--saida", str(out),
        "--no-paginas-pdf",
    ])
    assert proc.returncode == 0, (
        f"textura_apendice failed ({proc.returncode})\n"
        f"STDOUT:\n{proc.stdout}\nSTDERR:\n{proc.stderr}"
    )
    assert out.is_file()
    return out


class TestNearGolden:
    @pytest.mark.local_perf
    @pytest.mark.skipif(
        bool(os.environ.get("CI") or os.environ.get("GITHUB_ACTIONS")),
        reason="wall-clock baseline is machine-local; skip in CI",
    )
    def test_runtime_within_phase0_budget(self, near_bundle: dict):
        """Same-machine ±20% budget vs committed desktop baseline.

        Not a cross-machine gate. Re-run locally: pytest -m local_perf.
        """
        assert BASELINE.is_file(), f"missing timing baseline {BASELINE}"
        ref = float(json.loads(BASELINE.read_text(encoding="utf-8"))[
            "near_wall_seconds"])
        elapsed = float(near_bundle["seconds"])
        limit = max(ref * 1.20, ref + 2.0)
        assert elapsed <= limit, (
            f"near wall-clock {elapsed:.3f}s exceeds local Phase-0 budget "
            f"(baseline {ref:.3f}s, limit {limit:.3f}s)"
        )

    def test_exit_and_sheets(self, near_xlsx: Path):
        sheets = pd.ExcelFile(near_xlsx).sheet_names
        assert sheets == SHEETS_NEAR

    def test_concordancia_columns(self, near_xlsx: Path):
        df = pd.read_excel(near_xlsx, sheet_name="8_Concordancia")
        assert list(df.columns) == COLS_CONCORDANCIA
        assert "dominio_janela" in df.columns
        assert "revisao_sugerida" in df.columns

    def test_nuclear_counts(self, near_xlsx: Path):
        df = pd.read_excel(near_xlsx, sheet_name="8_Concordancia")
        nuc = _truthy_nuclear(df["nuclear"])
        assert len(df) == N_HITS
        assert int(nuc.sum()) == N_NUCLEAR
        assert int((~nuc).sum()) == N_NON_NUCLEAR

    def test_motivo_on_non_nuclear(self, near_xlsx: Path):
        df = pd.read_excel(near_xlsx, sheet_name="8_Concordancia")
        non = df.loc[~_truthy_nuclear(df["nuclear"])]
        vazios = non["motivo_exclusao"].fillna("").astype(str).str.strip().eq("")
        assert int(vazios.sum()) == 0, non.loc[
            vazios, ["matched_form", "motivo_exclusao"]
        ].to_string()

    def test_duplicate_tag_families(self, near_xlsx: Path):
        df = pd.read_excel(near_xlsx, sheet_name="8_Concordancia")
        blob = " | ".join(df["candidato_duplicado"].fillna("").astype(str))
        assert RX_C.search(blob), "missing citacao_entre_doc_ids:C####"
        assert RX_P.search(blob), "missing passagem_sobreposta:P####"
        assert RX_J.search(blob), "missing janela_sobreposta:J####"
        # stable first-group ids on this fixture
        assert "C0001" in blob
        assert "P0001" in blob
        assert "J0001" in blob

    def test_signals_present(self, near_xlsx: Path):
        df = pd.read_excel(near_xlsx, sheet_name="8_Concordancia")
        rev = df["revisao_sugerida"].fillna("").astype(str)
        assert rev.str.contains("coordenacao_heterogenea").any()
        assert rev.str.contains("associativa_com_nao_textural").any(), (
            "associative heterog. flag missing — wrapper may have been dropped"
        )
        dom = df["dominio_janela"].fillna("").astype(str)
        assert (dom == "geologia").any()
        # di-uniform extracted (campo includes *uniform)
        forms = df["matched_form"].astype(str).str.lower()
        assert forms.eq("di-uniform").any()


class TestAnaliseGolden:
    def test_associacao_sheet_and_or(self, analise_xlsx: Path):
        sheets = pd.ExcelFile(analise_xlsx).sheet_names
        assert "9_Associacao" in sheets
        assert "1_Resumo" in sheets
        assoc = pd.read_excel(
            analise_xlsx, sheet_name="9_Associacao", header=3)
        assert "razao_possib" in assoc.columns
        assert "canonical_term" in assoc.columns
        sub = assoc.loc[
            assoc["canonical_term"].astype(str) == "homogeneous",
            "razao_possib",
        ]
        assert len(sub) == 1
        assert float(sub.iloc[0]) == pytest.approx(OR_HOMOGENEOUS, rel=1e-3)


class TestApendiceGolden:
    def test_docx_table_rows(self, apendice_docx: Path):
        doc = Document(str(apendice_docx))
        assert doc.tables, "expected at least one table"
        data_rows = sum(max(0, len(t.rows) - 1) for t in doc.tables)
        assert data_rows == N_APENDICE_DATA_ROWS
