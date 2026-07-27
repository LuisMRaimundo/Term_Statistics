#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
textura_apendice.py — Fase 3: apêndice de concordância (DOCX)
=============================================================

Projecção legível das atribuições nucleares (genuínas) a partir da
folha ``8_Concordancia`` do Excel da fase 1 (revisto) ou da fase 2.

O documento segue o contrato tipográfico de
``Apendice_Concordancia_tabela_v2.docx``:

  • título «Apêndice [N] — Concordância das atribuições genuínas, por termo»
  • parágrafo introdutório com contagens
  • secção Heading 2 por termo/padrão: ``uniform* (n = 25)``
  • tabela de 2 colunas: Excerto (janela de contexto) | Fonte

Uso:
    python textura_apendice.py --xlsx resultado_near.xlsx
    python textura_apendice.py --xlsx resultado_analise.xlsx \\
        --refs obras_subjacentes.csv --saida Apendice_Concordancia.docx

Títulos APA corruptos (fallback do filename)? Gere antes o catálogo:
    python textura_apa7.py --xlsx resultado_near.xlsx --saida refs_apa7.xlsx
    python textura_apendice.py --xlsx resultado_near.xlsx --refs refs_apa7.xlsx
(A localização de páginas no PDF é independente — ``--paginas-pdf``.)

Gera sempre dois ficheiros:
  • Apendice_Concordancia.docx          — tipografia de publicação
  • Apendice_Concordancia_links.docx    — igual, com hiperligações na Fonte

    python textura_apendice.py --reformat "concordância tabelas_Anexo.docx"
"""

from __future__ import annotations

import argparse
import os
import re
import shutil
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import quote

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Constantes (espelham o apêndice de referência)
# Larguras em TWIPS (dxa). NÃO usar int(Twips(...)) no XML — isso é EMU.
# ---------------------------------------------------------------------------

COL_EXCERTO = "Excerto (janela de contexto)"
COL_FONTE = "Fonte"
# half-points (w:sz): 19 = 9,5 pt; 23 = 11,5 pt; 26 = 13 pt
SZ_CORPO = 19
SZ_H2 = 23
SZ_TITULO = 26
# grelha do Apendice_Concordancia_tabela_v2.docx
DXA_TABELA = 9026
DXA_COL_EXCERTO = 4381
DXA_COL_FONTE = 4645

_REF_CANDIDATOS = (
    Path(__file__).resolve().parent / "Apendice_Concordancia_tabela_v2.docx",
    Path(r"C:\Users\lmr20\Desktop\Tesaurus e Dicionários\CLASSES TEXTURAIS"
         r"\UNIFORM\Uniform_usos de domínio\Apendice_Concordancia_tabela_v2.docx"),
)

COLS_PADRAO = ("query_pattern", "canonical_term", "termo_tipo")
COLS_CONTEXTO = ("contexto",)
COLS_FONTE_DIRECTA = (
    "fonte_apa", "apa7", "fonte", "referencia", "referencia_apa",
)
COLS_CAMINHO = ("caminho_ficheiro", "caminho", "doc_id", "url")
COLS_PAGINA = ("pagina", "page", "p", "n")

_RX_MD_ITALIC = re.compile(r"\*([^*]+)\*")
_RX_EXT = re.compile(
    r"\.(pdf|txt|docx?|html?|xhtml|rtf|odt|epub|xml|tei)$", re.I)


# ---------------------------------------------------------------------------
# Leitura
# ---------------------------------------------------------------------------

def _truthy_nuclear(val) -> bool:
    if isinstance(val, bool):
        return val
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return False
    s = str(val).strip().lower()
    return s in {"1", "true", "verdadeiro", "yes", "sim", "t", "y"}


def ler_concordancia(xlsx: Path, folha: str | None = None) -> pd.DataFrame:
    """Lê ``8_Concordancia``, tolerando o cabeçalho meta da fase 2."""
    with pd.ExcelFile(xlsx) as xl:
        sheets = list(xl.sheet_names)
        if folha:
            if folha not in sheets:
                raise SystemExit(
                    f"Folha '{folha}' inexistente. Disponíveis: {sheets}")
            nome = folha
        elif "8_Concordancia" in sheets:
            nome = "8_Concordancia"
        else:
            raise SystemExit(
                "Falta a folha 8_Concordancia. "
                "Corra a fase 1 (textura_near / textura_search) primeiro.")
        df = pd.read_excel(xl, sheet_name=nome)
        if "contexto" not in df.columns:
            # fase 2: meta (unidade/N) nas primeiras linhas; dados em header=3
            df = pd.read_excel(xl, sheet_name=nome, header=3)
        if "contexto" not in df.columns:
            raise SystemExit(
                f"A folha '{nome}' não tem coluna 'contexto'. "
                f"Colunas: {list(df.columns)}")
    df = df.dropna(how="all")
    df = df[df["contexto"].notna()].copy()
    df["contexto"] = df["contexto"].astype(str)
    return df.reset_index(drop=True)


def filtrar_genuinas(df: pd.DataFrame, *, incluir_nao_nucleares: bool = False
                     ) -> pd.DataFrame:
    """Retém atribuições genuínas/nucleares (contrato do apêndice)."""
    out = df.copy()
    if "atribuicao" in out.columns and not incluir_nao_nucleares:
        atr = out["atribuicao"].astype(str).str.strip().str.lower()
        # legado spaCy: genuína / incidental
        if atr.isin({"genuína", "genuina", "incidental"}).any():
            out = out.loc[atr.isin({"genuína", "genuina"})].copy()
    if "nuclear" in out.columns and not incluir_nao_nucleares:
        out = out.loc[out["nuclear"].map(_truthy_nuclear)].copy()
    return out.reset_index(drop=True)


def escolher_coluna_agrupamento(df: pd.DataFrame, preferida: str | None) -> str:
    if preferida:
        if preferida not in df.columns:
            raise SystemExit(
                f"Coluna de agrupamento '{preferida}' ausente. "
                f"Disponíveis: {list(df.columns)}")
        return preferida
    for c in COLS_PADRAO:
        if c in df.columns and df[c].notna().any():
            return c
    raise SystemExit(
        "Não há coluna de agrupamento "
        "(query_pattern / canonical_term / termo_tipo).")


# ---------------------------------------------------------------------------
# Catálogo de referências APA
# ---------------------------------------------------------------------------

def _norm_chave(s: str) -> str:
    s = str(s or "").strip().lower().replace("\\", "/")
    s = _RX_EXT.sub("", s)
    return Path(s).name


def carregar_refs(caminho: Path | None) -> dict[str, str]:
    """Mapa chave normalizada → APA7.

    Aceita CSV/TSV/Excel com colunas ``apa7`` (ou ``fonte``/``referencia``)
    e chaves em ``underlying_file``, ``caminho``, ``doc_id``, ``ficheiro``, etc.
    """
    if caminho is None:
        return {}
    if not caminho.exists():
        raise SystemExit(f"Catálogo de referências não encontrado: {caminho}")
    if caminho.suffix.lower() in {".xlsx", ".xlsm"}:
        raw = pd.read_excel(caminho)
    else:
        # detecção simples de separador
        amostra = caminho.read_text(encoding="utf-8", errors="replace")[:2048]
        sep = ";" if amostra.count(";") > amostra.count(",") else ","
        if "\t" in amostra and amostra.count("\t") > amostra.count(sep):
            sep = "\t"
        raw = pd.read_csv(caminho, sep=sep)

    col_apa = next((c for c in ("apa7", "fonte_apa", "fonte", "referencia",
                                "referencia_apa", "apa")
                    if c in raw.columns), None)
    if col_apa is None:
        raise SystemExit(
            f"Catálogo sem coluna APA (apa7/fonte/referencia): {caminho}")

    chaves = [c for c in (
        "underlying_file", "ficheiro", "arquivo", "caminho_ficheiro",
        "caminho", "file", "filename", "doc_id", "id_doc", "binder_document",
    ) if c in raw.columns]
    if not chaves:
        raise SystemExit(
            f"Catálogo sem coluna de chave (doc_id/caminho/underlying_file): "
            f"{caminho}")

    mapa: dict[str, str] = {}
    for _, row in raw.iterrows():
        apa = row.get(col_apa)
        if apa is None or (isinstance(apa, float) and pd.isna(apa)):
            continue
        apa_s = str(apa).strip()
        if not apa_s or apa_s.startswith("{") or apa_s.lower() == "nan":
            continue
        for ck in chaves:
            val = row.get(ck)
            if val is None or (isinstance(val, float) and pd.isna(val)):
                continue
            mapa[_norm_chave(str(val))] = apa_s
    return mapa


def _primeira_col(row: pd.Series, cols: tuple[str, ...]):
    for c in cols:
        if c in row.index:
            v = row[c]
            if v is None or (isinstance(v, float) and pd.isna(v)):
                continue
            s = str(v).strip()
            if s and s.lower() != "nan":
                return s
    return None


def _pagina_de(row: pd.Series) -> str | None:
    for c in COLS_PAGINA:
        if c not in row.index:
            continue
        v = row[c]
        if v is None or (isinstance(v, float) and pd.isna(v)):
            continue
        # 'n' na matriz KWIC pode ser contagem; só usar se parecer página
        if c == "n":
            try:
                n = int(float(v))
            except (TypeError, ValueError):
                continue
            if n <= 0 or n > 9999:
                continue
            return str(n)
        s = str(v).strip()
        if not s or s.lower() in {"nan", "none", "—", "-"}:
            continue
        # "p. 12" / "12" / "12-13"
        m = re.search(r"(\d+(?:\s*[–\-]\s*\d+)?)", s)
        return m.group(1).replace(" ", "") if m else s
    return None


# Palavras curtas completas — não cortar só por comprimento.
_SHORT_OK = frozenset("""
a an the or and of in on to for by as is it at no not so if we he she be do
am are was were its his her their our your this that these those with from
into onto over under than then also only both such per via vs etc
um uma o a os as de do da dos das em no na nos nas e ou que se
""".split())

_RX_ELLIPSIS = re.compile(r"(?:\.{3}|…)+")


@dataclass(frozen=True)
class PaginaRef:
    """Página bibliográfica vs. página do ficheiro PDF (não confundir).

    ``impressa`` — número/etiqueta impressa no artigo/livro (APA: p. N),
    quando recuperável (coluna Excel ou page labels do PDF).
    ``pdf`` — índice 1-based no ficheiro PDF (folha do documento digital).
    """
    impressa: str | None = None
    pdf: int | None = None

    def cite(self) -> str | None:
        """Texto a anexar à fonte: ``(p. 45)`` ou ``(PDF p. 12)``."""
        if self.impressa:
            return f"(p. {self.impressa})"
        if self.pdf is not None:
            return f"(PDF p. {self.pdf})"
        return None


def _desembrulhar_excerto(texto: str) -> str:
    """Remove aspas/guillemets e reticências exteriores do excerto."""
    t = str(texto or "").strip()
    t = t.replace("«", "").replace("»", "").replace("“", "").replace("”", "")
    if (t.startswith('"') and t.endswith('"')) or (
            t.startswith("'") and t.endswith("'")):
        t = t[1:-1].strip()
    t = _RX_ELLIPSIS.sub("…", t).strip()
    while t.startswith("…"):
        t = t[1:].lstrip(" \t-–—")
    while t.endswith("…"):
        t = t[:-1].rstrip(" \t-–—")
    return t.strip()


def _termos_protegidos(*formas: str | None) -> list[str]:
    """Formas lexicais que a limpeza KWIC nunca pode remover."""
    out: list[str] = []
    vistos: set[str] = set()
    for f in formas:
        if f is None or (isinstance(f, float) and pd.isna(f)):
            continue
        s = str(f).strip()
        if not s or s.lower() in {"nan", "none"}:
            continue
        # multiword → guardar a sequência e cada token ≥ 2
        for peça in (s, *s.split()):
            p = peça.strip().casefold()
            if len(p) >= 2 and p not in vistos:
                vistos.add(p)
                out.append(p)
    return out


def _contem_termo_protegido(texto: str, protegidos: list[str]) -> bool:
    if not protegidos:
        return False
    t = f" {texto.casefold()} "
    for p in protegidos:
        if f" {p} " in t or t.strip() == p:
            return True
        # prefixo com hífen (di-uniform, etc.)
        if re.search(rf"(?<![\w]){re.escape(p)}", texto, flags=re.I):
            return True
    return False


def limpar_truncatura_kwic(
        texto: str,
        *,
        proteger: list[str] | tuple[str, ...] | None = None) -> str:
    """Remove caudas KWIC cortadas a meio (ex.: «, in ora»).

    Nunca remove os termos de pesquisa / nó (`proteger`). Se um corte
    proposto apagasse `matched_form`, `no`, etc., o corte é recusado.
    """
    t = " ".join(str(texto or "").split())
    if not t:
        return t
    protegidos = list(proteger or [])

    def _seguro(candidato: str) -> bool:
        """Aceita o corte só se todos os protegidos ainda presentes
        (ou se já faltavam no original — não restauramos)."""
        if not protegidos:
            return True
        for p in protegidos:
            # só exigir preservação se o original já continha o termo
            if _contem_termo_protegido(t, [p]) and not _contem_termo_protegido(
                    candidato, [p]):
                return False
        return True

    # 1) cauda curta após vírgula / ; / :
    if not re.search(r'[.!?]["\']?\s*$', t):
        cortes = list(re.finditer(r"[,;:]\s+", t))
        if cortes:
            ult = cortes[-1]
            cauda = t[ult.end():].strip()
            toks = cauda.split()
            if 1 <= len(toks) <= 3 and not re.search(r"[.!?]$", cauda):
                # não cortar se a cauda contém termo protegido
                if not _contem_termo_protegido(cauda, protegidos):
                    cand = t[: ult.start()].rstrip()
                    if _seguro(cand):
                        t = cand

    def _ultima_palavra(tok: str) -> str:
        tok = _RX_ELLIPSIS.sub("", tok)
        return re.sub(r"^[^\w]+|[^\w]+$", "", tok, flags=re.UNICODE)

    # 2) cotos finais curtos (≤3), nunca termos protegidos
    parts = t.split()
    if parts and not re.search(r"[.!?]$", parts[-1]):
        ultima = _ultima_palavra(parts[-1])
        u = ultima.casefold()
        # NÃO usar p.startswith(u): «texture».startswith(«t») bloquearia cotos.
        if (ultima and len(u) <= 3 and u not in _SHORT_OK
                and u not in protegidos):
            cand = " ".join(parts[:-1])
            if _seguro(cand):
                t = cand

    # 3) partículas gramaticais finais (of / the / in …) — máx. 3
    if not re.search(r'[.!?]["\']?\s*$', t):
        parts = t.split()
        removidos = 0
        while (parts and removidos < 3
               and not re.search(r"[.!?]$", parts[-1])):
            ultima = _ultima_palavra(parts[-1])
            u = ultima.casefold()
            if not ultima or u not in _SHORT_OK:
                break
            if u in protegidos:
                break
            cand = " ".join(parts[:-1])
            if not _seguro(cand):
                break
            parts = parts[:-1]
            t = cand
            removidos += 1

    return t.rstrip(" ,;:")


def caminho_local_de_fonte(url_ou_caminho: str | None) -> Path | None:
    """Resolve ``file:///…`` ou caminho Windows para Path local existente."""
    if not url_ou_caminho:
        return None
    s = str(url_ou_caminho).strip()
    if not s or s.lower() in {"nan", "none"}:
        return None
    if s.lower().startswith("file:"):
        from urllib.parse import unquote, urlparse
        u = urlparse(s)
        path = unquote(u.path or "")
        # file:///E:/foo → /E:/foo no Windows
        if re.match(r"^/[A-Za-z]:", path):
            path = path[1:]
        path = path.replace("/", "\\") if os_name_is_windows() else path
        p = Path(path)
        return p if p.is_file() else None
    p = Path(s)
    if p.is_file():
        return p
    return None


def os_name_is_windows() -> bool:
    return sys.platform.startswith("win")


def _sonda_pesquisa(snippet: str) -> list[str]:
    """Gera sondas de pesquisa cada vez mais curtas a partir do miolo."""
    t = limpar_truncatura_kwic(_desembrulhar_excerto(snippet))
    t = re.sub(r"\s+", " ", t).strip()
    if len(t) < 12:
        return []
    sondas = []
    # preferir o terço central (menos afectado por truncatura)
    n = len(t)
    if n >= 48:
        i0, i1 = n // 4, (3 * n) // 4
        sondas.append(t[i0:i1].strip())
    sondas.append(t[:80].strip())
    sondas.append(t[-80:].strip() if n > 80 else t)
    # sondas de ~40 chars em volta de palavras longas
    palavras = [w for w in re.findall(r"[A-Za-zÀ-ÿ]{5,}", t)]
    if len(palavras) >= 3:
        sondas.append(" ".join(palavras[len(palavras)//2 - 1:
                                        len(palavras)//2 + 2]))
    # únicos, por comprimento descrescente
    vistos = set()
    out = []
    for s in sorted((x for x in sondas if len(x) >= 12),
                    key=len, reverse=True):
        k = s.casefold()
        if k not in vistos:
            vistos.add(k)
            out.append(s)
    return out


def _chave_snippet(snippet: str) -> str:
    return re.sub(r"\s+", " ", snippet or "").strip().casefold()[:180]


def _label_impressa(page, pdf_index_1based: int, *, tem_page_labels: bool
                    ) -> str | None:
    """Etiqueta impressa do PDF, se distinta da mera folha digital."""
    try:
        lab = (page.get_label() or "").strip()
    except Exception:
        lab = ""
    if not lab:
        return None
    # Sem /PageLabels no PDF, get_label() costuma espelhar 1..n — não é
    # página de artigo/livro; nesses casos devolvemos None.
    if not tem_page_labels and lab == str(pdf_index_1based):
        return None
    return lab


def _resolver_snippets_num_pdf(caminho_str: str,
                               snippets: list[str]) -> list[PaginaRef]:
    """Abre UM pdf e resolve todos os snippets (uma passagem)."""
    caminho = Path(caminho_str)
    vazios = [PaginaRef() for _ in snippets]
    if not caminho.is_file() or caminho.suffix.lower() != ".pdf":
        return vazios
    try:
        import fitz
    except ImportError:
        return vazios
    try:
        doc = fitz.open(caminho)
    except Exception:
        return vazios

    out: list[PaginaRef] = []
    try:
        try:
            tem_labels = bool(doc.get_page_labels())
        except Exception:
            tem_labels = False
        # texto normalizado por página (1×) — fallback rápido
        textos = []
        for i in range(doc.page_count):
            try:
                textos.append(re.sub(
                    r"\s+", " ", doc.load_page(i).get_text("text")).casefold())
            except Exception:
                textos.append("")

        for snippet in snippets:
            ref = PaginaRef()
            idx0: int | None = None
            for sonda in _sonda_pesquisa(snippet):
                for i in range(doc.page_count):
                    page = doc.load_page(i)
                    try:
                        hits = page.search_for(sonda)
                    except Exception:
                        hits = []
                    if not hits:
                        sonda_n = re.sub(r"\s+", " ", sonda).casefold()
                        if not (sonda_n and sonda_n in textos[i]):
                            continue
                    idx0 = i
                    break
                if idx0 is not None:
                    break
            if idx0 is not None:
                page = doc.load_page(idx0)
                pdf_n = idx0 + 1
                impressa = _label_impressa(
                    page, pdf_n, tem_page_labels=tem_labels)
                ref = PaginaRef(impressa=impressa, pdf=pdf_n)
            out.append(ref)
    finally:
        doc.close()
    return out if len(out) == len(snippets) else vazios


def precarregar_paginas_pdf(
        tarefas: list[tuple[str, str]],
        *,
        workers: int | None = None) -> dict[tuple[str, str], PaginaRef]:
    """Resolve páginas em paralelo: 1 tarefa = 1 PDF (todos os seus excertos).

    ``tarefas``: lista (caminho_pdf, snippet). Usa vários blocos de CPU via
    threads (I/O + pesquisa PyMuPDF); cada PDF abre-se uma só vez.
    """
    if not tarefas:
        return {}
    por_pdf: dict[str, list[str]] = {}
    ordem: dict[str, list[str]] = {}
    for caminho, snippet in tarefas:
        chave_s = _chave_snippet(snippet)
        por_pdf.setdefault(caminho, [])
        ordem.setdefault(caminho, [])
        if chave_s not in ordem[caminho]:
            ordem[caminho].append(chave_s)
            por_pdf[caminho].append(snippet)

    n_pdf = len(por_pdf)
    cpu = os.cpu_count() or 4
    n_workers = workers if workers and workers > 0 else min(n_pdf, max(4, cpu * 2))
    print(f"      PDF: {n_pdf} ficheiros | {len(tarefas)} excertos | "
          f"workers={n_workers}", flush=True)

    resultado: dict[tuple[str, str], PaginaRef] = {}

    def _job(item: tuple[str, list[str]]) -> tuple[str, list[str], list[PaginaRef]]:
        path, snips = item
        return path, [_chave_snippet(s) for s in snips], _resolver_snippets_num_pdf(
            path, snips)

    with ThreadPoolExecutor(max_workers=n_workers) as pool:
        futs = [pool.submit(_job, (p, por_pdf[p])) for p in por_pdf]
        feitos = 0
        for fut in as_completed(futs):
            path, chaves, refs = fut.result()
            for c, r in zip(chaves, refs):
                resultado[(path, c)] = r
            feitos += 1
            if feitos % 10 == 0 or feitos == n_pdf:
                print(f"      PDF progresso: {feitos}/{n_pdf}", flush=True)
    return resultado


def resolver_pagina(row: pd.Series, *,
                    contexto: str | None = None,
                    procurar_pdf: bool = True,
                    mapa_pdf: dict[tuple[str, str], PaginaRef] | None = None,
                    ) -> PaginaRef | None:
    """Página bibliográfica (Excel) ou ref. PDF pré-carregada."""
    pag = _pagina_de(row)
    if pag:
        # Coluna Excel = página de obra/artigo (não folha do PDF).
        return PaginaRef(impressa=pag, pdf=None)
    if not procurar_pdf:
        return None
    url = resolver_url(row)
    caminho = caminho_local_de_fonte(url)
    if caminho is None:
        caminho = caminho_local_de_fonte(
            _primeira_col(row, ("caminho_ficheiro", "caminho")))
    if caminho is None:
        return None
    ctx = contexto if contexto is not None else str(
        row.get("contexto", "") or "")
    ch = _chave_snippet(ctx)
    try:
        key_path = str(caminho.resolve())
    except OSError:
        key_path = str(caminho)
    if mapa_pdf is not None:
        for k in ((key_path, ch), (str(caminho), ch)):
            if k in mapa_pdf:
                return mapa_pdf[k]
    # fallback síncrono (um PDF) se a chave não estiver no mapa
    refs = _resolver_snippets_num_pdf(key_path, [ctx])
    return refs[0] if refs else None


def _fonte_fallback(caminho: str | None) -> str:
    if not caminho:
        return "Fonte não identificada. (n.d.)."
    stem = Path(str(caminho).replace("\\", "/")).stem
    rotulo = stem.replace("_", " ").strip() or stem
    return f"{rotulo}. (n.d.)."


def caminho_saida_links(saida: Path) -> Path:
    """``Apendice.docx`` → ``Apendice_links.docx``."""
    return saida.with_name(f"{saida.stem}_links{saida.suffix}")


def resolver_url(row: pd.Series) -> str | None:
    """URL clicável: coluna ``url``, senão ``file:///`` a partir do caminho."""
    for c in ("url", "URL", "link"):
        if c not in row.index:
            continue
        v = row[c]
        if v is None or (isinstance(v, float) and pd.isna(v)):
            continue
        s = str(v).strip()
        if not s or s.lower() in {"nan", "none", "nat"}:
            continue
        if "://" in s or s.lower().startswith("file:"):
            return s
        # URL relativa / DOI nu
        if s.startswith("10.") and "/" in s:
            return f"https://doi.org/{s}"
        if s.startswith("www."):
            return "https://" + s
    caminho = _primeira_col(row, ("caminho_ficheiro", "caminho"))
    if not caminho:
        return None
    p = Path(caminho)
    if not p.is_absolute():
        # manter como file URI mesmo relativo (Word resolve no SO)
        pass
    try:
        # Path.as_uri() exige absoluto no Windows
        if p.is_absolute():
            return p.as_uri()
    except (ValueError, OSError):
        pass
    # fallback manual
    norm = str(caminho).replace("\\", "/")
    if re.match(r"^[A-Za-z]:/", norm):
        return "file:///" + quote(norm, safe="/:()[]@!$&'()*+,;=.-~")
    return "file://" + quote(norm, safe="/:()[]@!$&'()*+,;=.-~")


def resolver_fonte(row: pd.Series, refs: dict[str, str],
                   *, procurar_pdf: bool = False,
                   mapa_pdf: dict[tuple[str, str], PaginaRef] | None = None,
                   pagina_ref: PaginaRef | None = None) -> str:
    """Referência APA / fallback — **sem** número de página.

    A página vai a seguir ao excerto (coluna esquerda), não na Fonte.
    """
    _ = (procurar_pdf, mapa_pdf, pagina_ref)  # API estável; página no excerto
    directa = _primeira_col(row, COLS_FONTE_DIRECTA)
    if directa:
        base = directa
    else:
        base = None
        for c in COLS_CAMINHO:
            if c not in row.index:
                continue
            v = row[c]
            if v is None or (isinstance(v, float) and pd.isna(v)):
                continue
            chave = _norm_chave(str(v))
            if chave in refs:
                base = refs[chave]
                break
            # tentativa pelo nome completo do path
            chave2 = _norm_chave(Path(str(v).replace("\\", "/")).name)
            if chave2 in refs:
                base = refs[chave2]
                break
        if base is None:
            base = _fonte_fallback(_primeira_col(row, COLS_CAMINHO))

    base = base.rstrip()
    # APA 7: never add a period after a DOI/URL terminal element.
    if re.search(r"https?://(?:dx\.)?doi\.org/\S+$", base, re.I):
        return base
    if re.search(r"https?://\S+$", base, re.I):
        return base
    if base.endswith("."):
        return base
    return base + "."


def formatar_excerto(
        contexto: str,
        *,
        matched_form: str | None = None,
        no: str | None = None,
        canonical_term: str | None = None,
        termo_forma: str | None = None,
        pagina_ref: PaginaRef | None = None) -> str:
    """Excerto tipográfico: ``\"… miolo…\"`` (+ página à direita do snipet).

    Os termos de pesquisa (`matched_form`, `no`, …) são protegidos: a
    limpeza de truncatura nunca os remove do miolo. A página (obra ou
    folha PDF) segue o snipet, nunca a coluna Fonte/APA.
    """
    bruto = _desembrulhar_excerto(contexto)
    if not bruto:
        snip = '"… …"'
    else:
        protegidos = _termos_protegidos(
            matched_form, termo_forma, no, canonical_term)
        miolo = limpar_truncatura_kwic(bruto, proteger=protegidos)
        if not miolo:
            miolo = bruto
        fecha = "" if re.search(r'[.!?]["\']?\s*$', miolo) else "…"
        snip = f'"… {miolo}{fecha}"'
    if pagina_ref is not None:
        cite = pagina_ref.cite()
        if cite:
            return f"{snip} {cite}"
    return snip


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------

def _set_run_sz(run, half_points: int, *, bold: bool | None = None,
                italic: bool | None = None):
    """Aplica w:sz / w:szCs em half-points (como o apêndice de referência)."""
    rpr = run._element.get_or_add_rPr()
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), str(half_points))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def _add_text_with_italic_md(paragraph, text: str, *, bold: bool = False,
                             italic: bool | None = None,
                             sz: int = SZ_CORPO):
    """Escreve texto; ``*itálico*`` (APA markdown) vira run itálico."""
    _clear_paragraph(paragraph)
    pos = 0
    for m in _RX_MD_ITALIC.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            _set_run_sz(run, sz, bold=bold or None, italic=italic)
        run = paragraph.add_run(m.group(1))
        _set_run_sz(run, sz, bold=bold or None, italic=True)
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_sz(run, sz, bold=bold or None, italic=italic)


def _set_dxa(el_parent, tag: str, dxa: int):
    el = el_parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        el_parent.append(el)
    el.set(qn("w:w"), str(dxa))
    el.set(qn("w:type"), "dxa")


def _set_cell_width(cell, dxa: int):
    tcPr = cell._tc.get_or_add_tcPr()
    _set_dxa(tcPr, "w:tcW", dxa)


def _set_table_layout(table):
    """Normal Table + grelha/bordas/margens do apêndice de referência."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # remover estilo Table Grid (parece «por formatar»); usar Normal Table
    stil = tblPr.find(qn("w:tblStyle"))
    if stil is not None:
        tblPr.remove(stil)
    try:
        table.style = "Normal Table"
    except KeyError:
        pass

    _set_dxa(tblPr, "w:tblW", DXA_TABELA)

    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")

    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblPr.append(mar)
    for lado, val in (("left", "10"), ("right", "10")):
        el = mar.find(qn(f"w:{lado}"))
        if el is None:
            el = OxmlElement(f"w:{lado}")
            mar.append(el)
        el.set(qn("w:w"), val)
        el.set(qn("w:type"), "dxa")

    # grelha
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for dxa in (DXA_COL_EXCERTO, DXA_COL_FONTE):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(dxa))
        grid.append(gc)


def _preencher_celula(cell, texto: str, *, bold: bool = False,
                      italic: bool | None = None, dxa: int):
    _set_cell_width(cell, dxa)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_text_with_italic_md(p, texto, bold=bold, italic=italic, sz=SZ_CORPO)


def _add_hyperlink(paragraph, texto: str, url: str, *, sz: int = SZ_CORPO):
    """Insere hiperligação externa (estilo azul sublinhado)."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # estilo Hyperlink do Word, se existir
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    cor = OxmlElement("w:color")
    cor.set(qn("w:val"), "0563C1")
    rPr.append(cor)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(sz))
        rPr.append(el)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = texto
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _preencher_celula_fonte(cell, texto: str, url: str | None, *, dxa: int):
    """Fonte com hiperligação; sem URL marca «[sem URL]» para revisão rápida."""
    _set_cell_width(cell, dxa)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _clear_paragraph(p)
    if url:
        _add_hyperlink(p, texto, url, sz=SZ_CORPO)
    else:
        _add_text_with_italic_md(p, texto, sz=SZ_CORPO)
        run = p.add_run(" [sem URL]")
        _set_run_sz(run, SZ_CORPO, bold=True)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def _introducao(*, n_atr: int, n_obras: int, apendice: str,
                cf_secao: str, anexo_dados: str) -> str:
    return (
        f"O presente apêndice documenta, em formato de concordância, as "
        f"{n_atr} atribuições genuínas do levantamento de usos de domínio "
        f"(cf. §{cf_secao}), distribuídas por {n_obras} obras distintas e "
        f"organizadas alfabeticamente por termo pesquisado. A coluna "
        f"esquerda reproduz a janela de contexto indexada tipograficamente "
        f"normalizada (aspas rectas e reticências); cotos finais de KWIC "
        f"claramente truncados são omitidos, sem remover os termos de "
        f"pesquisa. A coluna direita apresenta a referência da fonte "
        f"(APA, 7.ª ed., quando disponível no catálogo), sem número de página. "
        f"A página de obra/artigo (p. N) ou, na sua falta, a folha do "
        f"ficheiro digital (PDF p. N) indica-se a seguir ao excerto, na "
        f"coluna esquerda. As entradas sem metadados "
        f"descritivos recuperáveis identificam-se pelo nome estável do "
        f"ficheiro. Os critérios de extracção e validação, a classificação "
        f"integral de cada instância e os testes estatísticos constam do "
        f"anexo de dados digital [{anexo_dados}], de que este apêndice "
        f"(Apêndice [{apendice}]) é a projecção legível."
    )


def _limpar_corpo(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _documento_base(template: Path | None = None) -> Document:
    """Parte do molde de referência (estilos/secção) ou A4 com margens iguais."""
    caminho = template
    if caminho is None:
        caminho = next((p for p in _REF_CANDIDATOS if p.exists()), None)
    if caminho is not None and caminho.exists():
        doc = Document(str(caminho))
        _limpar_corpo(doc)
        return doc
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    return doc


def _escrever_titulo_intro(doc: Document, *, titulo: str, intro: str) -> None:
    tit = doc.add_paragraph()
    tit.paragraph_format.space_after = Pt(4)
    run = tit.add_run(titulo)
    _set_run_sz(run, SZ_TITULO, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    _add_text_with_italic_md(p, intro, italic=True, sz=SZ_CORPO)


def _escrever_secao_tabela(
        doc: Document, rotulo: str,
        linhas: list[tuple[str, str, str | None]],
        *, com_hiperligacoes: bool = False) -> None:
    h = doc.add_heading(f"{rotulo} (n = {len(linhas)})", level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(5)
    for run in h.runs:
        _set_run_sz(run, SZ_H2, bold=True)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _set_table_layout(table)
    _preencher_celula(table.rows[0].cells[0], COL_EXCERTO,
                      bold=True, dxa=DXA_COL_EXCERTO)
    cab_fonte = COL_FONTE + (" (com hiperligação)" if com_hiperligacoes else "")
    _preencher_celula(table.rows[0].cells[1], cab_fonte,
                      bold=True, dxa=DXA_COL_FONTE)
    for excerto, fonte, url in linhas:
        row = table.add_row().cells
        _preencher_celula(row[0], excerto, dxa=DXA_COL_EXCERTO)
        if com_hiperligacoes:
            _preencher_celula_fonte(row[1], fonte, url, dxa=DXA_COL_FONTE)
        else:
            _preencher_celula(row[1], fonte, dxa=DXA_COL_FONTE)


def _mapa_paginas_para_df(
        df: pd.DataFrame, *, workers: int | None = None
) -> dict[tuple[str, str], PaginaRef]:
    """Recolhe tarefas PDF e resolve em paralelo."""
    tarefas: list[tuple[str, str]] = []
    for _, r in df.iterrows():
        if _pagina_de(r):
            continue  # já há página de obra no Excel
        url = resolver_url(r)
        caminho = caminho_local_de_fonte(url)
        if caminho is None:
            caminho = caminho_local_de_fonte(
                _primeira_col(r, ("caminho_ficheiro", "caminho")))
        if caminho is None or caminho.suffix.lower() != ".pdf":
            continue
        try:
            key_path = str(caminho.resolve())
        except OSError:
            key_path = str(caminho)
        tarefas.append((key_path, str(r.get("contexto", "") or "")))
    return precarregar_paginas_pdf(tarefas, workers=workers)


def construir_documento(
    df: pd.DataFrame,
    *,
    col_grupo: str,
    refs: dict[str, str],
    apendice: str = "X",
    cf_secao: str = "7.1.2",
    anexo_dados: str = "X",
    template: Path | None = None,
    com_hiperligacoes: bool = False,
    procurar_pdf: bool = True,
    pdf_workers: int | None = None,
    mapa_pdf: dict[tuple[str, str], PaginaRef] | None = None,
) -> Document:
    doc = _documento_base(template)

    col_obra = next((c for c in ("doc_id", "caminho_ficheiro", "caminho")
                     if c in df.columns), None)
    n_atr = len(df)
    n_obras = int(df[col_obra].nunique()) if col_obra else 0

    titulo = (f"Apêndice [{apendice}] — Concordância das atribuições "
              f"genuínas, por termo")
    if com_hiperligacoes:
        titulo += " [versão com hiperligações]"
    intro = _introducao(n_atr=n_atr, n_obras=n_obras, apendice=apendice,
                        cf_secao=cf_secao, anexo_dados=anexo_dados)
    if com_hiperligacoes:
        intro += (' Nesta versão de verificação, a coluna Fonte liga ao '
                  'ficheiro ou URL de origem; entradas sem ligação '
                  'aparecem marcadas com "[sem URL]".')

    _escrever_titulo_intro(doc, titulo=titulo, intro=intro)

    df = df.copy()
    df["_grupo"] = df[col_grupo].astype(str).str.strip()
    df = df[df["_grupo"].ne("") & df["_grupo"].str.lower().ne("nan")]
    grupos = sorted(df["_grupo"].unique(), key=lambda s: s.casefold())

    if procurar_pdf and mapa_pdf is None:
        mapa_pdf = _mapa_paginas_para_df(df, workers=pdf_workers)

    n_obra = n_pdf = n_label = 0
    for g in grupos:
        sub = df.loc[df["_grupo"] == g]
        if col_obra:
            sub = sub.sort_values(
                by=[c for c in (col_obra, "contexto") if c in sub.columns],
                kind="mergesort",
            )
        linhas = []
        for _, r in sub.iterrows():
            pref = resolver_pagina(
                r, contexto=str(r.get("contexto", "") or ""),
                procurar_pdf=procurar_pdf, mapa_pdf=mapa_pdf,
            )
            ex = formatar_excerto(
                r["contexto"],
                matched_form=r.get("matched_form"),
                no=r.get("no"),
                canonical_term=r.get("canonical_term"),
                termo_forma=r.get("termo_forma"),
                pagina_ref=pref,
            )
            fo = resolver_fonte(r, refs)
            if pref:
                if pref.impressa:
                    n_obra += 1
                    if pref.pdf is not None:
                        n_label += 1
                elif pref.pdf is not None:
                    n_pdf += 1
            linhas.append((ex, fo, resolver_url(r)))
        _escrever_secao_tabela(doc, g, linhas,
                               com_hiperligacoes=com_hiperligacoes)

    if procurar_pdf:
        print(
            f"      páginas: obra/artigo (p. N)={n_obra} "
            f"| só folha PDF (PDF p. N)={n_pdf} "
            f"| com page-label PDF={n_label} | total linhas={n_atr}",
            flush=True,
        )
    return doc


def extrair_blocos_docx(docx: Path) -> tuple[str, str, list[tuple[str, list[tuple[str, str]]]]]:
    """Lê um apêndice já gerado → (título, intro, [(termo, [(excerto, fonte)])])."""
    d = Document(str(docx))
    titulo = ""
    intro = ""
    blocos: list[tuple[str, list[tuple[str, str]]]] = []
    headings = [p for p in d.paragraphs
                if p.style and str(p.style.name).startswith("Heading")
                and p.text.strip()]
    # título / intro = primeiros Normal não vazios
    for p in d.paragraphs:
        if p.style and str(p.style.name).startswith("Heading"):
            break
        t = p.text.strip()
        if not t:
            continue
        if not titulo:
            titulo = t
        elif not intro:
            intro = t
    if len(headings) != len(d.tables):
        raise SystemExit(
            f"Estrutura inesperada em {docx.name}: "
            f"{len(headings)} headings vs {len(d.tables)} tabelas.")
    rx = re.compile(r"^(.*?)\s*\(\s*n\s*=\s*\d+\s*\)\s*$", re.I)
    for h, table in zip(headings, d.tables):
        m = rx.match(h.text.strip())
        termo = m.group(1).strip() if m else h.text.strip()
        linhas: list[tuple[str, str]] = []
        for row in table.rows[1:]:
            cells = row.cells
            if len(cells) < 2:
                continue
            ex = cells[0].text.strip()
            fo = cells[1].text.strip()
            if not ex and not fo:
                continue
            linhas.append((ex, fo))
        blocos.append((termo, linhas))
    return titulo, intro, blocos


def reformatar_docx(origem: Path, saida: Path | None = None,
                    *, template: Path | None = None,
                    backup: bool = True) -> Path:
    """Regrava um apêndice sem formatação com o molde tipográfico correcto."""
    dest = saida or origem
    titulo, intro, blocos = extrair_blocos_docx(origem)
    if not blocos:
        raise SystemExit(f"Nenhuma secção encontrada em {origem}")
    if backup and dest.resolve() == origem.resolve():
        bak = origem.with_suffix(origem.suffix + ".bak")
        if not bak.exists():
            shutil.copy2(origem, bak)
    doc = _documento_base(template)
    _escrever_titulo_intro(doc, titulo=titulo or (
        "Apêndice [X] — Concordância das atribuições genuínas, por termo"),
        intro=intro or "")
    for termo, linhas in blocos:
        # reformat: sem URLs (só tipografia)
        _escrever_secao_tabela(
            doc, termo, [(ex, fo, None) for ex, fo in linhas],
            com_hiperligacoes=False)
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


def gerar_apendice(
    xlsx: Path,
    saida: Path,
    *,
    folha: str | None = None,
    refs_path: Path | None = None,
    agrupar: str | None = None,
    apendice: str = "X",
    cf_secao: str = "7.1.2",
    anexo_dados: str = "X",
    incluir_nao_nucleares: bool = False,
    template: Path | None = None,
    procurar_pdf: bool = True,
    pdf_workers: int | None = None,
) -> dict:
    """Gera os dois DOCX (publicação + versão com hiperligações)."""
    print(f"  fonte Excel: {xlsx}", flush=True)
    brutas = ler_concordancia(xlsx, folha=folha)
    n_false = 0
    if "nuclear" in brutas.columns and not incluir_nao_nucleares:
        n_false = int((~brutas["nuclear"].map(_truthy_nuclear)).sum())
    genuinas = filtrar_genuinas(
        brutas, incluir_nao_nucleares=incluir_nao_nucleares)
    print(
        f"  filtro nuclear=TRUE: {len(brutas)} linhas lidas -> "
        f"{len(genuinas)} no apêndice"
        + (f" (ignoradas nuclear=FALSE: {n_false})" if n_false else ""),
        flush=True,
    )
    if genuinas.empty:
        raise SystemExit(
            "Nenhuma atribuição genuína/nuclear para documentar. "
            "Verifique a coluna 'nuclear' (ou use --incluir-nao-nucleares).")
    col_grupo = escolher_coluna_agrupamento(genuinas, agrupar)
    refs = carregar_refs(refs_path)
    saida.parent.mkdir(parents=True, exist_ok=True)
    saida_links = caminho_saida_links(saida)

    def _guardar(doc_obj: Document, caminho: Path) -> Path:
        try:
            doc_obj.save(str(caminho))
            return caminho
        except PermissionError:
            alt = caminho.with_name(
                f"{caminho.stem}_novo{caminho.suffix}")
            doc_obj.save(str(alt))
            print(
                f"AVISO: «{caminho.name}» está aberto/bloqueado; "
                f"gravei em {alt.name}",
                flush=True,
            )
            return alt

    mapa_pdf = None
    if procurar_pdf:
        try:
            import fitz  # noqa: F401
        except ImportError:
            print(
                "AVISO: --paginas-pdf pediu pesquisa no PDF, mas pymupdf "
                "não está instalado (pip install pymupdf). "
                "Continuo só com páginas já indexadas no Excel.",
                flush=True,
            )
            procurar_pdf = False
        else:
            # Uma só passagem paralela — reutilizada nos dois DOCX.
            mapa_pdf = _mapa_paginas_para_df(
                genuinas, workers=pdf_workers)

    doc = construir_documento(
        genuinas, col_grupo=col_grupo, refs=refs,
        apendice=apendice, cf_secao=cf_secao, anexo_dados=anexo_dados,
        template=template, com_hiperligacoes=False,
        procurar_pdf=procurar_pdf, pdf_workers=pdf_workers,
        mapa_pdf=mapa_pdf,
    )
    saida = _guardar(doc, saida)

    doc_l = construir_documento(
        genuinas, col_grupo=col_grupo, refs=refs,
        apendice=apendice, cf_secao=cf_secao, anexo_dados=anexo_dados,
        template=template, com_hiperligacoes=True,
        procurar_pdf=procurar_pdf, pdf_workers=pdf_workers,
        mapa_pdf=mapa_pdf,
    )
    saida_links = _guardar(doc_l, saida_links)

    n_com_url = sum(1 for _, r in genuinas.iterrows() if resolver_url(r))
    col_obra = next((c for c in ("doc_id", "caminho_ficheiro", "caminho")
                     if c in genuinas.columns), None)
    resumo = {
        "saida": str(saida),
        "saida_links": str(saida_links),
        "linhas_brutas": int(len(brutas)),
        "atribuições_genuínas": int(len(genuinas)),
        "com_url": int(n_com_url),
        "sem_url": int(len(genuinas) - n_com_url),
        "obras": int(genuinas[col_obra].nunique()) if col_obra else 0,
        "termos": int(genuinas[col_grupo].nunique()),
        "coluna_grupo": col_grupo,
        "refs_catalogadas": int(len(refs)),
    }
    return resumo


def _configurar_consola() -> None:
    for stream in (sys.stdout, sys.stderr):
        try:
            stream.reconfigure(encoding="utf-8", errors="replace")
        except Exception:
            pass


def main(argv: list[str] | None = None) -> int:
    _configurar_consola()
    ap = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    ap.add_argument("--xlsx", type=Path, default=None,
                    help="Excel fase 1 (revisto) ou fase 2 (com 8_Concordancia)")
    ap.add_argument("--reformat", type=Path, default=None,
                    help="reformatar um DOCX de apêndice já gerado (in-place "
                         "ou com --saida); cria .bak na primeira vez")
    ap.add_argument("--saida", type=Path, default=None,
                    help="DOCX de saída")
    ap.add_argument("--template", type=Path, default=None,
                    help="DOCX molde (omissão: Apendice_Concordancia_tabela_v2)")
    ap.add_argument("--folha", default=None,
                    help="folha de concordância (omissão: 8_Concordancia)")
    ap.add_argument("--refs", type=Path, default=None,
                    help="catálogo APA7 (CSV/TSV/Excel: apa7 + doc_id/ficheiro)")
    ap.add_argument("--agrupar", default=None,
                    choices=["query_pattern", "canonical_term", "termo_tipo"],
                    help="coluna de agrupamento (omissão: primeira disponível)")
    ap.add_argument("--apendice", default="X",
                    help="número/letra do apêndice no título (omissão: X)")
    ap.add_argument("--cf-secao", default="7.1.2",
                    help="secção referida no intro (cf. §…)")
    ap.add_argument("--anexo-dados", default="X",
                    help="rótulo do anexo de dados digital no intro")
    ap.add_argument("--incluir-nao-nucleares", action="store_true",
                    help="não filtrar por nuclear/atribuicao (diagnóstico)")
    ap.add_argument(
        "--paginas-pdf",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="localizar página no PDF ligado (omissão: ligado; "
             "desligar com --no-paginas-pdf). (p. N)=obra/artigo via "
             "page-label; (PDF p. N)=folha do ficheiro. Requer pymupdf.",
    )
    ap.add_argument("--pdf-workers", type=int, default=0,
                    help="nº de PDFs em paralelo (omissão: ~2×CPU, máx. "
                         "nº de ficheiros)")
    args = ap.parse_args(argv)

    if args.reformat is not None:
        if not args.reformat.exists():
            print(f"Ficheiro não encontrado: {args.reformat}", file=sys.stderr)
            return 2
        print(f"Fase 3 · reformatar ← {args.reformat.name}", flush=True)
        dest = reformatar_docx(
            args.reformat, args.saida, template=args.template)
        titulo, intro, blocos = extrair_blocos_docx(dest)
        n = sum(len(linhas) for _, linhas in blocos)
        print(f"  secções={len(blocos)} | linhas={n}", flush=True)
        print(f"\n=== FASE 3 (reformat) concluída ===\n  {dest}")
        return 0

    if args.xlsx is None:
        ap.error("indique --xlsx ou --reformat")
    if not args.xlsx.exists():
        print(f"Ficheiro não encontrado: {args.xlsx}", file=sys.stderr)
        return 2
    saida = args.saida or args.xlsx.with_name("Apendice_Concordancia.docx")
    print(f"Fase 3 · apêndice DOCX ← {args.xlsx.name}", flush=True)
    resumo = gerar_apendice(
        args.xlsx, saida,
        folha=args.folha,
        refs_path=args.refs,
        agrupar=args.agrupar,
        apendice=args.apendice,
        cf_secao=args.cf_secao,
        anexo_dados=args.anexo_dados,
        incluir_nao_nucleares=args.incluir_nao_nucleares,
        template=args.template,
        procurar_pdf=args.paginas_pdf,
        pdf_workers=(args.pdf_workers or None),
    )
    print(
        f"  genuínas={resumo['atribuições_genuínas']} "
        f"(de {resumo['linhas_brutas']} brutas) | "
        f"obras={resumo['obras']} | termos={resumo['termos']} "
        f"(grupo={resumo['coluna_grupo']}) | "
        f"refs={resumo['refs_catalogadas']} | "
        f"urls={resumo['com_url']}/{resumo['atribuições_genuínas']}",
        flush=True,
    )
    print(f"\n=== FASE 3 concluída ===\n  {resumo['saida']}"
          f"\n  {resumo['saida_links']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
